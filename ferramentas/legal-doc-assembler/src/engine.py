"""
Document Engine - Core rendering logic for legal document generation.

Uses docxtpl with Jinja2 and custom normalization filters.
Fault-tolerant: undefined variables remain visible in output.
"""

import json
from pathlib import Path
from typing import Any, Dict, List, Optional

import jinja2
from jinja2 import DebugUndefined
from docxtpl import DocxTemplate
from rich.console import Console
from rich.table import Table

from .normalizers import (
    normalize_name,
    normalize_address,
    format_cpf,
    format_cnpj,
    format_cep,
    format_oab,
    normalize_whitespace,
    normalize_punctuation,
    normalize_all,
)

console = Console()


class DocumentEngine:
    """
    Document rendering engine for legal documents.

    Features:
        - Jinja2 templating with docxtpl
        - Fault-tolerant: undefined variables show {{ var_name }} in output
        - Custom filters for Brazilian document formatting
        - Automatic text normalization

    Usage:
        engine = DocumentEngine()
        engine.render(
            template_path="template.docx",
            data={"nome": "João da Silva", "cpf": "12345678901"},
            output_path="output.docx"
        )
    """

    def __init__(self, auto_normalize: bool = True):
        """
        Initialize the document engine.

        Args:
            auto_normalize: If True, automatically apply text normalization
                to string values in data dict.
        """
        self.auto_normalize = auto_normalize
        self._setup_jinja_env()

    def _setup_jinja_env(self) -> jinja2.Environment:
        """
        Configure Jinja2 environment with:
            - DebugUndefined: keeps {{ var }} visible if undefined
            - Custom filters for Brazilian formatting
        """
        self.jinja_env = jinja2.Environment(
            undefined=DebugUndefined,
            autoescape=False,  # Don't escape for docx
        )

        # Register custom filters
        self.jinja_env.filters['nome'] = normalize_name
        self.jinja_env.filters['endereco'] = normalize_address
        self.jinja_env.filters['cpf'] = format_cpf
        self.jinja_env.filters['cnpj'] = format_cnpj
        self.jinja_env.filters['cep'] = format_cep
        self.jinja_env.filters['oab'] = format_oab
        self.jinja_env.filters['texto'] = lambda x: normalize_punctuation(
            normalize_whitespace(x)
        )

        return self.jinja_env

    def _preprocess_data(
        self,
        data: Dict[str, Any],
        field_types: Optional[Dict[str, str]] = None
    ) -> Dict[str, Any]:
        """
        Preprocess data before rendering.

        If auto_normalize is True and field_types not provided,
        applies basic whitespace normalization to all string values.

        If field_types is provided, uses normalize_all() for type-specific
        normalization.

        Args:
            data: Raw data dictionary
            field_types: Optional dict mapping field names to types
                ('name', 'address', 'cpf', 'cnpj', 'cep', 'oab', 'text', 'raw')

        Returns:
            Preprocessed data dictionary
        """
        if not self.auto_normalize:
            return data

        if field_types:
            return normalize_all(data, field_types)

        # Default: apply whitespace normalization to all strings
        result = {}
        for key, value in data.items():
            if isinstance(value, str):
                result[key] = normalize_whitespace(value)
            elif isinstance(value, dict):
                result[key] = self._preprocess_data(value)
            elif isinstance(value, list):
                result[key] = [
                    self._preprocess_data(item) if isinstance(item, dict)
                    else normalize_whitespace(item) if isinstance(item, str)
                    else item
                    for item in value
                ]
            else:
                result[key] = value

        return result

    def render(
        self,
        template_path: str | Path,
        data: Dict[str, Any],
        output_path: str | Path,
        field_types: Optional[Dict[str, str]] = None,
    ) -> Path:
        """
        Render a document from template and data.

        Args:
            template_path: Path to .docx template file
            data: Dictionary with template variables
            output_path: Path for output .docx file
            field_types: Optional dict mapping field names to normalization types

        Returns:
            Path to the generated document

        Raises:
            FileNotFoundError: If template doesn't exist
            ValueError: If template is invalid
        """
        template_path = Path(template_path)
        output_path = Path(output_path)

        # Validate template exists
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        # Preprocess data
        processed_data = self._preprocess_data(data, field_types)

        # Load template with custom Jinja environment
        try:
            doc = DocxTemplate(template_path)
            doc.render(processed_data, self.jinja_env)
        except Exception as e:
            raise ValueError(f"Error rendering template: {e}")

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Save document
        doc.save(output_path)

        console.print(f"[green]✔[/green] Document saved: {output_path}")

        return output_path

    def render_from_json(
        self,
        template_path: str | Path,
        json_path: str | Path,
        output_path: str | Path,
        field_types: Optional[Dict[str, str]] = None,
    ) -> Path:
        """
        Render document from template and JSON data file.

        Args:
            template_path: Path to .docx template
            json_path: Path to JSON file with data
            output_path: Path for output document
            field_types: Optional normalization type mapping

        Returns:
            Path to generated document
        """
        json_path = Path(json_path)

        if not json_path.exists():
            raise FileNotFoundError(f"JSON file not found: {json_path}")

        with open(json_path, 'r', encoding='utf-8') as f:
            data = json.load(f)

        return self.render(template_path, data, output_path, field_types)

    def get_template_variables(self, template_path: str | Path) -> List[str]:
        """
        Extract variable names from a template.

        Useful for understanding what data a template expects.

        Args:
            template_path: Path to .docx template

        Returns:
            List of variable names found in template
        """
        template_path = Path(template_path)

        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        doc = DocxTemplate(template_path)
        # Use our Jinja2 environment which has custom filters registered
        variables = doc.get_undeclared_template_variables(self.jinja_env)

        return sorted(list(variables))

    def get_template_text(self, template_path: str | Path) -> List[Dict[str, Any]]:
        """
        Extract text content from a template as a list of paragraphs.

        Each paragraph includes its text, index, and whether it contains
        Jinja2 variables.

        Args:
            template_path: Path to .docx template

        Returns:
            List of dicts with keys: 'index', 'text', 'has_variables'
        """
        from docx import Document
        import re

        template_path = Path(template_path)

        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        doc = Document(template_path)
        paragraphs = []

        jinja_pattern = re.compile(r'\{\{.*?\}\}')

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if text:  # Skip empty paragraphs
                paragraphs.append({
                    'index': i,
                    'text': text,
                    'has_variables': bool(jinja_pattern.search(text))
                })

        return paragraphs

    def mark_text_as_variable(
        self,
        template_path: str | Path,
        text_to_replace: str,
        variable_name: str,
        output_path: str | Path,
        filter_name: Optional[str] = None,
        replace_all: bool = True
    ) -> Path:
        """
        Replace specific text in template with Jinja2 variable placeholder.

        Args:
            template_path: Path to source .docx template
            text_to_replace: Exact text to find and replace
            variable_name: Name for the Jinja2 variable
            output_path: Path to save modified template
            filter_name: Optional filter to apply (nome, cpf, etc.)
            replace_all: If True, replace all occurrences; if False, only first

        Returns:
            Path to the modified template

        Example:
            engine.mark_text_as_variable(
                "template.docx",
                "MARIA DA SILVA",
                "nome",
                "marked_template.docx",
                filter_name="nome"
            )
            # Result: "MARIA DA SILVA" becomes "{{ nome|nome }}"
        """
        from docx import Document

        template_path = Path(template_path)
        output_path = Path(output_path)

        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

        doc = Document(template_path)

        # Build the variable placeholder
        if filter_name:
            placeholder = f"{{{{ {variable_name}|{filter_name} }}}}"
        else:
            placeholder = f"{{{{ {variable_name} }}}}"

        # Process each paragraph
        for para in doc.paragraphs:
            if text_to_replace in para.text:
                # Handle runs (formatted segments within paragraph)
                full_text = para.text
                if replace_all:
                    new_text = full_text.replace(text_to_replace, placeholder)
                else:
                    new_text = full_text.replace(text_to_replace, placeholder, 1)

                # Clear existing runs and add new text
                # Note: This may lose formatting - acceptable for MVP
                for run in para.runs:
                    run.text = ""
                if para.runs:
                    para.runs[0].text = new_text
                else:
                    para.add_run(new_text)

        # Ensure output directory exists
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        return output_path

    def find_markable_patterns(self, text: str) -> List[Dict[str, Any]]:
        """
        Find common patterns in text that could be marked as variables.

        Detects:
            - CPF (formatted or unformatted)
            - CNPJ (formatted or unformatted)
            - CEP
            - Dates (DD/MM/YYYY format)

        Args:
            text: Text to analyze

        Returns:
            List of dicts with: 'text', 'type', 'start', 'end', 'suggested_var', 'suggested_filter'
        """
        import re

        patterns = []

        # CPF patterns (formatted and unformatted)
        cpf_pattern = re.compile(r'\b\d{3}[.\s]?\d{3}[.\s]?\d{3}[-.\s]?\d{2}\b')
        for match in cpf_pattern.finditer(text):
            patterns.append({
                'text': match.group(),
                'type': 'cpf',
                'start': match.start(),
                'end': match.end(),
                'suggested_var': 'cpf',
                'suggested_filter': 'cpf'
            })

        # CNPJ patterns
        cnpj_pattern = re.compile(r'\b\d{2}[.\s]?\d{3}[.\s]?\d{3}[/\s]?\d{4}[-.\s]?\d{2}\b')
        for match in cnpj_pattern.finditer(text):
            patterns.append({
                'text': match.group(),
                'type': 'cnpj',
                'start': match.start(),
                'end': match.end(),
                'suggested_var': 'cnpj',
                'suggested_filter': 'cnpj'
            })

        # CEP patterns
        cep_pattern = re.compile(r'\b\d{5}[-.\s]?\d{3}\b')
        for match in cep_pattern.finditer(text):
            # Exclude if already matched as CPF/CNPJ
            if not any(p['start'] <= match.start() < p['end'] for p in patterns):
                patterns.append({
                    'text': match.group(),
                    'type': 'cep',
                    'start': match.start(),
                    'end': match.end(),
                    'suggested_var': 'cep',
                    'suggested_filter': 'cep'
                })

        # Date patterns (DD/MM/YYYY)
        date_pattern = re.compile(r'\b\d{1,2}/\d{1,2}/\d{2,4}\b')
        for match in date_pattern.finditer(text):
            patterns.append({
                'text': match.group(),
                'type': 'date',
                'start': match.start(),
                'end': match.end(),
                'suggested_var': 'data',
                'suggested_filter': None
            })

        return patterns

    def validate_data(
        self,
        template_path: str | Path,
        data: Dict[str, Any]
    ) -> Dict[str, List[str]]:
        """
        Validate data against template requirements.

        Args:
            template_path: Path to .docx template
            data: Data dictionary to validate

        Returns:
            Dict with 'missing' and 'extra' keys listing field names
        """
        required = set(self.get_template_variables(template_path))
        provided = set(data.keys())

        return {
            'missing': sorted(list(required - provided)),
            'extra': sorted(list(provided - required)),
        }

    def print_validation_report(
        self,
        template_path: str | Path,
        data: Dict[str, Any]
    ) -> None:
        """Print a formatted validation report to console."""
        validation = self.validate_data(template_path, data)

        table = Table(title="Data Validation Report")
        table.add_column("Status", style="bold")
        table.add_column("Fields")

        if validation['missing']:
            table.add_row(
                "[yellow]Missing[/yellow]",
                ", ".join(validation['missing'])
            )
        else:
            table.add_row("[green]Missing[/green]", "None")

        if validation['extra']:
            table.add_row(
                "[blue]Extra[/blue]",
                ", ".join(validation['extra'])
            )
        else:
            table.add_row("[green]Extra[/green]", "None")

        console.print(table)

        if validation['missing']:
            console.print(
                "\n[yellow]⚠ Warning:[/yellow] Missing fields will appear "
                "as {{ field_name }} in output."
            )
