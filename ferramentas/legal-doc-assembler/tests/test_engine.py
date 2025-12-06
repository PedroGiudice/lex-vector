# tests/test_engine.py
"""Integration tests for DocumentEngine."""

import json
import tempfile
from pathlib import Path

import pytest
from docx import Document

from src.engine import DocumentEngine


@pytest.fixture
def engine():
    """Create a DocumentEngine instance."""
    return DocumentEngine(auto_normalize=True)


@pytest.fixture
def engine_no_normalize():
    """Create a DocumentEngine instance without auto-normalization."""
    return DocumentEngine(auto_normalize=False)


@pytest.fixture
def temp_dir():
    """Create a temporary directory for test outputs."""
    with tempfile.TemporaryDirectory() as tmpdir:
        yield Path(tmpdir)


@pytest.fixture
def sample_template(temp_dir):
    """Create a simple test template."""
    from docx import Document as DocxDoc

    template_path = temp_dir / "test_template.docx"
    doc = DocxDoc()
    doc.add_paragraph("Nome: {{ nome }}")
    doc.add_paragraph("CPF: {{ cpf }}")
    doc.add_paragraph("RG: {{ rg }}")
    doc.add_paragraph("Endereço: {{ endereco }}")
    doc.save(template_path)

    return template_path


class TestDocumentEngineInit:
    """Tests for DocumentEngine initialization."""

    def test_init_default(self):
        """Test default initialization."""
        engine = DocumentEngine()
        assert engine.auto_normalize is True
        assert engine.jinja_env is not None

    def test_init_no_normalize(self):
        """Test initialization without auto-normalization."""
        engine = DocumentEngine(auto_normalize=False)
        assert engine.auto_normalize is False

    def test_jinja_filters_registered(self):
        """Test that all custom filters are registered."""
        engine = DocumentEngine()
        filters = engine.jinja_env.filters

        assert 'nome' in filters
        assert 'endereco' in filters
        assert 'cpf' in filters
        assert 'cnpj' in filters
        assert 'cep' in filters
        assert 'oab' in filters
        assert 'texto' in filters


class TestDocumentEngineRender:
    """Tests for DocumentEngine render method."""

    def test_render_basic(self, engine, sample_template, temp_dir):
        """Test basic document rendering."""
        data = {
            "nome": "João da Silva",
            "cpf": "12345678901",
            "rg": "12.345.678-9",
            "endereco": "Rua das Flores, 123",
        }

        output_path = temp_dir / "output.docx"
        result = engine.render(sample_template, data, output_path)

        assert result.exists()
        assert result.stat().st_size > 0

        # Verify content
        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "João da Silva" in text
        assert "12.345.678-9" in text

    def test_render_missing_field_preserved(self, engine, sample_template, temp_dir):
        """Test that missing fields show {{ field_name }} in output."""
        data = {
            "nome": "João da Silva",
            "cpf": "12345678901",
            # RG intentionally omitted
            # endereco intentionally omitted
        }

        output_path = temp_dir / "output_missing.docx"
        result = engine.render(sample_template, data, output_path)

        assert result.exists()

        # Verify missing fields are preserved as {{ field }}
        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "{{ rg }}" in text
        assert "{{ endereco }}" in text

    def test_render_with_normalization(self, engine, sample_template, temp_dir):
        """Test automatic whitespace normalization."""
        data = {
            "nome": "  MARIA   DA  SILVA  ",  # Extra spaces
            "cpf": "12345678901",
            "rg": "123456789",
            "endereco": "R.  DAS  FLORES",
        }

        output_path = temp_dir / "output_norm.docx"
        result = engine.render(sample_template, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        # Check that extra whitespace was normalized
        assert "  " not in text.split("Nome:")[1].split("\n")[0]

    def test_render_without_normalization(self, engine_no_normalize, sample_template, temp_dir):
        """Test rendering without auto-normalization."""
        data = {
            "nome": "  MARIA   DA  SILVA  ",  # Extra spaces preserved
            "cpf": "12345678901",
            "rg": "123456789",
            "endereco": "R.  DAS  FLORES",
        }

        output_path = temp_dir / "output_no_norm.docx"
        result = engine_no_normalize.render(sample_template, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        # Check that extra whitespace was preserved (no normalization)
        assert "MARIA   DA  SILVA" in text

    def test_render_template_not_found(self, engine, temp_dir):
        """Test error when template doesn't exist."""
        with pytest.raises(FileNotFoundError, match="Template not found"):
            engine.render(
                temp_dir / "nonexistent.docx",
                {"nome": "Test"},
                temp_dir / "output.docx"
            )

    def test_render_creates_output_directory(self, engine, sample_template, temp_dir):
        """Test that output directory is created if it doesn't exist."""
        data = {"nome": "Test", "cpf": "12345678901", "rg": "123", "endereco": "Rua"}
        output_path = temp_dir / "new_dir" / "subdir" / "output.docx"

        result = engine.render(sample_template, data, output_path)

        assert result.exists()
        assert result.parent.exists()


class TestDocumentEngineRenderFromJson:
    """Tests for DocumentEngine render_from_json method."""

    def test_render_from_json(self, engine, sample_template, temp_dir):
        """Test rendering from JSON file."""
        json_path = temp_dir / "data.json"
        data = {
            "nome": "João",
            "cpf": "12345678901",
            "rg": "123456789",
            "endereco": "Rua Teste",
        }

        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False)

        output_path = temp_dir / "output_json.docx"
        result = engine.render_from_json(sample_template, json_path, output_path)

        assert result.exists()

    def test_render_from_json_file_not_found(self, engine, sample_template, temp_dir):
        """Test error when JSON file doesn't exist."""
        with pytest.raises(FileNotFoundError, match="JSON file not found"):
            engine.render_from_json(
                sample_template,
                temp_dir / "nonexistent.json",
                temp_dir / "output.docx"
            )


class TestDocumentEngineTemplateVariables:
    """Tests for template variable extraction."""

    def test_get_template_variables(self, engine, sample_template):
        """Test extraction of template variables."""
        variables = engine.get_template_variables(sample_template)

        assert "nome" in variables
        assert "cpf" in variables
        assert "rg" in variables
        assert "endereco" in variables

    def test_get_template_variables_sorted(self, engine, sample_template):
        """Test that variables are returned sorted."""
        variables = engine.get_template_variables(sample_template)

        assert variables == sorted(variables)

    def test_get_template_variables_not_found(self, engine, temp_dir):
        """Test error when template doesn't exist."""
        with pytest.raises(FileNotFoundError, match="Template not found"):
            engine.get_template_variables(temp_dir / "nonexistent.docx")


class TestDocumentEngineValidation:
    """Tests for data validation."""

    def test_validate_data_all_present(self, engine, sample_template):
        """Test validation when all fields are present."""
        data = {
            "nome": "João",
            "cpf": "123",
            "rg": "456",
            "endereco": "Rua",
        }

        validation = engine.validate_data(sample_template, data)

        assert validation["missing"] == []
        assert validation["extra"] == []

    def test_validate_data_missing_fields(self, engine, sample_template):
        """Test validation with missing fields."""
        data = {
            "nome": "João",
            # cpf missing
            # rg missing
            "endereco": "Rua",
        }

        validation = engine.validate_data(sample_template, data)

        assert "cpf" in validation["missing"]
        assert "rg" in validation["missing"]
        assert validation["extra"] == []

    def test_validate_data_extra_fields(self, engine, sample_template):
        """Test validation with extra fields."""
        data = {
            "nome": "João",
            "cpf": "123",
            "rg": "456",
            "endereco": "Rua",
            "extra_field": "value",
            "another_extra": "value2",
        }

        validation = engine.validate_data(sample_template, data)

        assert validation["missing"] == []
        assert "extra_field" in validation["extra"]
        assert "another_extra" in validation["extra"]

    def test_validate_data_mixed(self, engine, sample_template):
        """Test validation with both missing and extra fields."""
        data = {
            "nome": "João",
            # cpf missing
            "rg": "456",
            "endereco": "Rua",
            "extra_field": "value",
        }

        validation = engine.validate_data(sample_template, data)

        assert "cpf" in validation["missing"]
        assert "extra_field" in validation["extra"]


class TestDocumentEngineFilters:
    """Tests for Jinja2 custom filters in templates."""

    def test_cpf_filter(self, engine, temp_dir):
        """Test CPF filter in template."""
        from docx import Document as DocxDoc

        template_path = temp_dir / "cpf_template.docx"
        doc = DocxDoc()
        doc.add_paragraph("CPF: {{ cpf_raw|cpf }}")
        doc.save(template_path)

        data = {"cpf_raw": "12345678901"}

        output_path = temp_dir / "output_cpf.docx"
        result = engine.render(template_path, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "123.456.789-01" in text

    def test_nome_filter(self, engine, temp_dir):
        """Test nome filter in template."""
        from docx import Document as DocxDoc

        template_path = temp_dir / "nome_template.docx"
        doc = DocxDoc()
        doc.add_paragraph("Nome: {{ nome_raw|nome }}")
        doc.save(template_path)

        data = {"nome_raw": "MARIA DA SILVA"}

        output_path = temp_dir / "output_nome.docx"
        result = engine.render(template_path, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Maria da Silva" in text

    def test_endereco_filter(self, engine, temp_dir):
        """Test endereco filter in template."""
        from docx import Document as DocxDoc

        template_path = temp_dir / "endereco_template.docx"
        doc = DocxDoc()
        doc.add_paragraph("Endereço: {{ end_raw|endereco }}")
        doc.save(template_path)

        data = {"end_raw": "AV. BRASIL N 500"}

        output_path = temp_dir / "output_endereco.docx"
        result = engine.render(template_path, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "Avenida" in text
        assert "nº" in text

    def test_cnpj_filter(self, engine, temp_dir):
        """Test CNPJ filter in template."""
        from docx import Document as DocxDoc

        template_path = temp_dir / "cnpj_template.docx"
        doc = DocxDoc()
        doc.add_paragraph("CNPJ: {{ cnpj_raw|cnpj }}")
        doc.save(template_path)

        data = {"cnpj_raw": "12345678000199"}

        output_path = temp_dir / "output_cnpj.docx"
        result = engine.render(template_path, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "12.345.678/0001-99" in text

    def test_cep_filter(self, engine, temp_dir):
        """Test CEP filter in template."""
        from docx import Document as DocxDoc

        template_path = temp_dir / "cep_template.docx"
        doc = DocxDoc()
        doc.add_paragraph("CEP: {{ cep_raw|cep }}")
        doc.save(template_path)

        data = {"cep_raw": "01310100"}

        output_path = temp_dir / "output_cep.docx"
        result = engine.render(template_path, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "01310-100" in text


class TestDocumentEngineUTF8:
    """Tests for UTF-8 character handling."""

    def test_utf8_characters(self, engine, temp_dir):
        """Test handling of UTF-8 Brazilian characters."""
        from docx import Document as DocxDoc

        template_path = temp_dir / "utf8_template.docx"
        doc = DocxDoc()
        doc.add_paragraph("Nome: {{ nome }}")
        doc.add_paragraph("Descrição: {{ descricao }}")
        doc.save(template_path)

        data = {
            "nome": "José Antônio da Conceição",
            "descricao": "Profissão: médico ortopedista - área de atuação"
        }

        output_path = temp_dir / "output_utf8.docx"
        result = engine.render(template_path, data, output_path)

        # Verify UTF-8 characters preserved
        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "José" in text
        assert "Antônio" in text
        assert "Conceição" in text
        assert "médico" in text
        assert "área" in text

    def test_utf8_accented_uppercase(self, engine, temp_dir):
        """Test uppercase accented characters."""
        from docx import Document as DocxDoc

        template_path = temp_dir / "utf8_upper_template.docx"
        doc = DocxDoc()
        doc.add_paragraph("{{ texto }}")
        doc.save(template_path)

        data = {"texto": "AÇÃO REVISIONAL - CLÁUSULA ABUSIVA"}

        output_path = temp_dir / "output_utf8_upper.docx"
        result = engine.render(template_path, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "AÇÃO" in text
        assert "CLÁUSULA" in text


class TestDocumentEngineNestedData:
    """Tests for nested data structures."""

    def test_nested_dict_normalization(self, engine, temp_dir):
        """Test normalization of nested dictionaries."""
        from docx import Document as DocxDoc

        template_path = temp_dir / "nested_template.docx"
        doc = DocxDoc()
        doc.add_paragraph("Nome: {{ pessoa.nome }}")
        doc.add_paragraph("Cidade: {{ pessoa.endereco.cidade }}")
        doc.save(template_path)

        data = {
            "pessoa": {
                "nome": "  João  ",
                "endereco": {
                    "cidade": "  São Paulo  "
                }
            }
        }

        output_path = temp_dir / "output_nested.docx"
        result = engine.render(template_path, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        # Check nested values were normalized
        assert "João" in text
        assert "  João  " not in text

    def test_list_normalization(self, engine, temp_dir):
        """Test normalization of lists."""
        from docx import Document as DocxDoc

        template_path = temp_dir / "list_template.docx"
        doc = DocxDoc()
        doc.add_paragraph("{% for item in items %}{{ item }} {% endfor %}")
        doc.save(template_path)

        data = {
            "items": ["  item1  ", "  item2  ", "  item3  "]
        }

        output_path = temp_dir / "output_list.docx"
        result = engine.render(template_path, data, output_path)

        doc = Document(result)
        text = "\n".join([p.text for p in doc.paragraphs])

        assert "item1" in text
        assert "  item1  " not in text


class TestDocumentEngineTextExtraction:
    """Tests for template text extraction."""

    def test_get_template_text_returns_paragraphs(self, temp_dir):
        """Test that get_template_text extracts paragraphs from docx."""
        from docx import Document as DocxDoc
        from src.engine import DocumentEngine

        # Create a test docx
        doc = DocxDoc()
        doc.add_paragraph("First paragraph with some text.")
        doc.add_paragraph("Second paragraph with more content.")
        doc.add_paragraph("Third paragraph.")
        test_path = temp_dir / "test_template.docx"
        doc.save(test_path)

        engine = DocumentEngine()
        paragraphs = engine.get_template_text(str(test_path))

        assert len(paragraphs) == 3
        assert paragraphs[0]["text"] == "First paragraph with some text."
        assert paragraphs[1]["text"] == "Second paragraph with more content."
        assert paragraphs[2]["text"] == "Third paragraph."
        assert all("index" in p for p in paragraphs)


class TestDocumentEngineMarkVariable:
    """Tests for marking text as Jinja2 variables."""

    def test_mark_text_as_variable(self, temp_dir):
        """Test marking specific text as a Jinja2 variable in docx."""
        from docx import Document as DocxDoc
        from src.engine import DocumentEngine

        # Create a test docx
        doc = DocxDoc()
        doc.add_paragraph("O cliente MARIA DA SILVA compareceu.")
        doc.add_paragraph("CPF: 12345678901")
        test_path = temp_dir / "test_mark.docx"
        doc.save(test_path)

        engine = DocumentEngine()

        # Mark "MARIA DA SILVA" as variable "nome"
        output_path = temp_dir / "marked.docx"
        result = engine.mark_text_as_variable(
            template_path=str(test_path),
            text_to_replace="MARIA DA SILVA",
            variable_name="nome",
            output_path=str(output_path)
        )

        # Verify the output contains the variable
        paragraphs = engine.get_template_text(str(output_path))
        assert "{{ nome }}" in paragraphs[0]["text"]
        assert paragraphs[0]["has_variables"] is True

    def test_mark_text_as_variable_with_filter(self, temp_dir):
        """Test marking text with a normalization filter."""
        from docx import Document as DocxDoc
        from src.engine import DocumentEngine

        doc = DocxDoc()
        doc.add_paragraph("CPF: 12345678901")
        test_path = temp_dir / "test_filter.docx"
        doc.save(test_path)

        engine = DocumentEngine()
        output_path = temp_dir / "filtered.docx"

        engine.mark_text_as_variable(
            template_path=str(test_path),
            text_to_replace="12345678901",
            variable_name="cpf",
            output_path=str(output_path),
            filter_name="cpf"
        )

        paragraphs = engine.get_template_text(str(output_path))
        assert "{{ cpf|cpf }}" in paragraphs[0]["text"]


class TestDocumentEnginePatternDetection:
    """Tests for automatic pattern detection in text."""

    def test_find_markable_patterns_detects_cpf(self):
        """Test that CPF patterns are detected."""
        from src.engine import DocumentEngine

        engine = DocumentEngine()
        text = "Cliente: MARIA DA SILVA, CPF: 123.456.789-01"
        patterns = engine.find_markable_patterns(text)

        cpf_patterns = [p for p in patterns if p["type"] == "cpf"]
        assert len(cpf_patterns) == 1
        assert cpf_patterns[0]["text"] == "123.456.789-01"
        assert cpf_patterns[0]["suggested_var"] == "cpf"
        assert cpf_patterns[0]["suggested_filter"] == "cpf"

    def test_find_markable_patterns_detects_cnpj(self):
        """Test that CNPJ patterns are detected."""
        from src.engine import DocumentEngine

        engine = DocumentEngine()
        text = "Empresa LTDA, CNPJ: 12.345.678/0001-99"
        patterns = engine.find_markable_patterns(text)

        cnpj_patterns = [p for p in patterns if p["type"] == "cnpj"]
        assert len(cnpj_patterns) == 1
        assert cnpj_patterns[0]["text"] == "12.345.678/0001-99"
        assert cnpj_patterns[0]["suggested_var"] == "cnpj"

    def test_find_markable_patterns_detects_cep(self):
        """Test that CEP patterns are detected."""
        from src.engine import DocumentEngine

        engine = DocumentEngine()
        text = "Endereco: Rua das Flores, 123 - CEP 01310-100"
        patterns = engine.find_markable_patterns(text)

        cep_patterns = [p for p in patterns if p["type"] == "cep"]
        assert len(cep_patterns) == 1
        assert cep_patterns[0]["text"] == "01310-100"

    def test_find_markable_patterns_detects_date(self):
        """Test that date patterns are detected."""
        from src.engine import DocumentEngine

        engine = DocumentEngine()
        text = "Sao Paulo, 15/06/2024"
        patterns = engine.find_markable_patterns(text)

        date_patterns = [p for p in patterns if p["type"] == "date"]
        assert len(date_patterns) == 1
        assert date_patterns[0]["text"] == "15/06/2024"
        assert date_patterns[0]["suggested_var"] == "data"

    def test_find_markable_patterns_multiple(self):
        """Test detection of multiple patterns in same text."""
        from src.engine import DocumentEngine

        engine = DocumentEngine()
        text = "CPF: 123.456.789-01, CNPJ: 12.345.678/0001-99, Data: 01/01/2025"
        patterns = engine.find_markable_patterns(text)

        assert len(patterns) >= 3
        types = [p["type"] for p in patterns]
        assert "cpf" in types
        assert "cnpj" in types
        assert "date" in types
