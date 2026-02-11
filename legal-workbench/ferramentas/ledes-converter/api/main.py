import os
import sys
import logging

# Add shared module path for logging and Sentry
sys.path.insert(0, '/app')

# Initialize Sentry BEFORE importing FastAPI for proper instrumentation
try:
    from shared.sentry_config import init_sentry
    init_sentry("ledes-converter")
except ImportError:
    pass  # Sentry not available, continue without it

# Configure structured JSON logging
from shared.logging_config import setup_logging
from shared.middleware import RequestIDMiddleware

log_level = getattr(logging, os.getenv("LOG_LEVEL", "INFO").upper(), logging.INFO)
logger = setup_logging("ledes-converter", level=log_level)

from fastapi import FastAPI, UploadFile, File, HTTPException, Request, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from typing import Annotated, Optional
from dataclasses import asdict
import docx
import tempfile
import io
import zipfile
from collections import defaultdict
import time
import magic
import json

from .models import (
    LedesData, ConversionResponse, HealthResponse, LineItem, LedesConfig,
    MatterRequest, MatterResponse, ValidationIssueResponse, ValidationResponse,
    BatchResultItem, BatchConversionResponse,
    StructuredConversionRequest, StructuredLineItemInput, TextConversionRequest,
)
from .matter_store import MatterStore, Matter
from .task_codes import classify_task_code, classify_activity_code
from .ledes_validator import validate_ledes_1998b
from .ledes_generator import (
    sanitize_string,
    sanitize_ledes_field,
    parse_currency,
    format_ledes_currency,
    format_date_ledes,
    extract_ledes_data,
    generate_ledes_1998b,
    MAX_DESCRIPTION_LENGTH,
    MAX_LINE_ITEMS,
)

app = FastAPI(
    title="LEDES Converter API",
    description="Convert DOCX invoices to LEDES 1998B format",
    version="1.0.0"
)

# Request ID middleware for request tracing
app.add_middleware(RequestIDMiddleware)

# Initialize matter store
matter_store = MatterStore()

# CORS middleware - configured for production security
allowed_origins = os.getenv(
    "ALLOWED_ORIGINS",
    "http://localhost,http://localhost:3000,https://tauri.localhost,http://tauri.localhost"
).split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST", "PUT", "DELETE"],
    allow_headers=["Content-Type", "Authorization"],
)

# Log startup
@app.on_event("startup")
async def startup_event():
    logger.info("Service starting", extra={"event": "startup", "version": "1.0.0"})

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_FILENAME_LENGTH = 255
ALLOWED_MIME_TYPES = [
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream"  # Some systems send DOCX as this
]

# Simple rate limiting (production should use Redis)
rate_limit_storage = defaultdict(list)


def rate_limit_check(client_ip: str, max_requests: int = 10, window_seconds: int = 60) -> bool:
    """Simple in-memory rate limiting. Production should use Redis."""
    now = time.time()
    # Clean old entries
    rate_limit_storage[client_ip] = [
        ts for ts in rate_limit_storage[client_ip]
        if now - ts < window_seconds
    ]

    if len(rate_limit_storage[client_ip]) >= max_requests:
        return False

    rate_limit_storage[client_ip].append(now)
    return True


def validate_file_type(content: bytes, filename: str) -> bool:
    """Validate file type using magic bytes (MIME detection)."""
    try:
        mime = magic.from_buffer(content, mime=True)

        # Check MIME type
        if mime not in ALLOWED_MIME_TYPES:
            logger.warning(f"Invalid MIME type detected: {mime} for file extension .docx")
            return False

        # Additional check: DOCX files should start with PK (ZIP header)
        if not content.startswith(b'PK'):
            logger.warning("File does not have valid ZIP/DOCX header")
            return False

        return True
    except Exception as e:
        logger.error(f"File type validation error: {e}")
        return False


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint for Docker/Kubernetes."""
    return HealthResponse(status="ok", service="ledes-converter")


@app.post("/convert/docx-to-ledes", response_model=ConversionResponse)
async def convert_docx_to_ledes(
    request: Request,
    file: Annotated[UploadFile, File(description="DOCX file to convert to LEDES")],
    config: Annotated[Optional[str], Form(description="JSON string with LEDES configuration")] = None
) -> ConversionResponse:
    """
    Convert a DOCX invoice file to LEDES 1998B format.

    - **file**: DOCX file containing invoice data (max 10MB)
    - **config**: Optional JSON string with LEDES configuration (law_firm_id, client_id, matter_id, etc.)
    - Returns: LEDES formatted content and extracted data

    Security features:
    - Rate limiting (10 requests/minute per IP)
    - MIME type validation using magic bytes
    - File size validation
    - Input sanitization
    """
    # Rate limiting
    client_ip = request.client.host if request.client else "unknown"
    if not rate_limit_check(client_ip):
        raise HTTPException(
            status_code=429,
            detail="Too many requests. Please try again later."
        )

    # Validate filename
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required.")

    if len(file.filename) > MAX_FILENAME_LENGTH:
        raise HTTPException(status_code=400, detail="Filename too long.")

    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a .docx file.")

    # Read and validate file size
    content = await file.read()
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")

    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB."
        )

    # Validate MIME type using magic bytes
    if not validate_file_type(content, file.filename):
        raise HTTPException(
            status_code=400,
            detail="Invalid file format. File does not appear to be a valid DOCX document."
        )

    tmp_path = ""
    try:
        # Save temp file with secure permissions
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", mode='wb') as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        # Set restrictive permissions
        os.chmod(tmp_path, 0o600)

        # Process DOCX with error handling
        try:
            doc = docx.Document(tmp_path)
        except Exception as docx_error:
            logger.error(f"Failed to parse DOCX: {docx_error}")
            raise HTTPException(
                status_code=400,
                detail="Unable to parse DOCX file. File may be corrupted or in an unsupported format."
            )

        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)

        text_content = "\n".join(full_text)

        # Validate extracted content
        if not text_content.strip():
            raise HTTPException(
                status_code=400,
                detail="No text content found in DOCX file."
            )

        # Extract Data
        extracted_data = extract_ledes_data(text_content)

        # Parse and validate config if provided
        ledes_config = None
        if config:
            try:
                config_dict = json.loads(config)

                # If matter_name is provided, load from store and merge
                if "matter_name" in config_dict and not config_dict.get("law_firm_id"):
                    matter = matter_store.get(config_dict["matter_name"])
                    if not matter:
                        raise HTTPException(
                            status_code=404,
                            detail=f"Matter not found: {config_dict['matter_name']}"
                        )
                    config_dict.setdefault("law_firm_id", matter.law_firm_id)
                    config_dict.setdefault("law_firm_name", matter.law_firm_name)
                    config_dict.setdefault("client_id", matter.client_id)
                    config_dict.setdefault("client_name", matter.client_name)
                    config_dict.setdefault("matter_id", matter.matter_id)
                    config_dict.setdefault("client_matter_id", matter.client_matter_id)
                    config_dict.setdefault("timekeeper_id", matter.timekeeper_id)
                    config_dict.setdefault("timekeeper_name", matter.timekeeper_name)
                    config_dict.setdefault("timekeeper_classification", matter.timekeeper_classification)
                    config_dict.setdefault("unit_cost", matter.unit_cost)
                    logger.info(f"Loaded matter config: {matter.matter_name}")

                ledes_config = LedesConfig(**config_dict)
                logger.info(f"Config provided: law_firm={ledes_config.law_firm_id}, client={ledes_config.client_id}, matter={ledes_config.matter_id}")
            except json.JSONDecodeError as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid JSON in config parameter: {str(e)}"
                )
            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid config data: {str(e)}"
                )

        # Merge config with extracted data (config takes precedence)
        if ledes_config:
            # Firm/Client/Matter
            extracted_data["law_firm_id"] = ledes_config.law_firm_id
            extracted_data["law_firm_name"] = ledes_config.law_firm_name
            extracted_data["client_id"] = ledes_config.client_id
            extracted_data["client_name"] = ledes_config.client_name
            extracted_data["matter_id"] = ledes_config.matter_id
            extracted_data["matter_name"] = ledes_config.matter_name
            extracted_data["client_matter_id"] = ledes_config.client_matter_id or ""
            # Timekeeper
            extracted_data["timekeeper_id"] = ledes_config.timekeeper_id or ""
            extracted_data["timekeeper_name"] = ledes_config.timekeeper_name or ""
            extracted_data["timekeeper_classification"] = ledes_config.timekeeper_classification or ""
            extracted_data["unit_cost"] = ledes_config.unit_cost or 0
            # Billing period
            extracted_data["billing_start_date"] = ledes_config.billing_start_date or ""
            extracted_data["billing_end_date"] = ledes_config.billing_end_date or ""
        else:
            # Backward compatibility: use placeholders if no config provided
            extracted_data["law_firm_id"] = ""
            extracted_data["client_matter_id"] = ""
            logger.warning("No config provided, using empty law_firm_id (backward compatibility mode)")

        # Validate extracted data
        if not extracted_data.get("invoice_number"):
            logger.warning("Invoice number not found in document")

        if not extracted_data.get("line_items"):
            raise HTTPException(
                status_code=400,
                detail="No line items found in invoice. Please check document format."
            )

        # Generate LEDES
        ledes_content = generate_ledes_1998b(extracted_data)

        logger.info(f"Successfully converted invoice (lines: {len(extracted_data['line_items'])})")

        # Create validated response using Pydantic models
        line_items = [LineItem(**item) for item in extracted_data["line_items"]]
        ledes_data = LedesData(
            invoice_date=extracted_data["invoice_date"],
            invoice_number=extracted_data["invoice_number"],
            client_id=extracted_data["client_id"],
            matter_id=extracted_data["matter_id"],
            invoice_total=extracted_data["invoice_total"],
            line_items=line_items
        )

        return ConversionResponse(
            filename=sanitize_string(file.filename, MAX_FILENAME_LENGTH),
            status="success",
            extracted_data=ledes_data,
            ledes_content=ledes_content
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Conversion failed: {type(e).__name__}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="Conversion failed due to an internal error. Please contact support if the issue persists."
        )
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception as cleanup_error:
                logger.error(f"Failed to cleanup temp file: {cleanup_error}")


@app.post("/convert/structured", response_model=ConversionResponse, tags=["Conversion"])
async def convert_structured(req: StructuredConversionRequest):
    """
    Convert structured form data directly to LEDES 1998B.

    Accepts pre-filled invoice data + line items from a form.
    Task/activity codes are auto-classified from descriptions if not provided.
    Matter config (firm, client, timekeeper) is loaded from the matter store.
    """
    matter = matter_store.get(req.matter_name)
    if not matter:
        raise HTTPException(status_code=404, detail=f"Matter not found: {req.matter_name}")

    line_items = []
    invoice_total = 0.0
    for item in req.line_items:
        task_code = item.task_code or classify_task_code(item.description)
        activity_code = item.activity_code or classify_activity_code(item.description)
        line_items.append({
            "description": item.description,
            "amount": item.amount,
            "task_code": task_code,
            "activity_code": activity_code,
        })
        invoice_total += item.amount

    data = {
        "invoice_date": req.invoice_date,
        "invoice_number": req.invoice_number,
        "invoice_total": invoice_total,
        "invoice_description": req.invoice_description,
        "billing_start_date": req.billing_start_date,
        "billing_end_date": req.billing_end_date,
        "client_id": matter.client_id,
        "client_name": matter.client_name,
        "matter_id": matter.matter_id,
        "matter_name": matter.matter_name,
        "client_matter_id": matter.client_matter_id,
        "law_firm_id": matter.law_firm_id,
        "law_firm_name": matter.law_firm_name,
        "timekeeper_id": matter.timekeeper_id,
        "timekeeper_name": matter.timekeeper_name,
        "timekeeper_classification": matter.timekeeper_classification,
        "unit_cost": matter.unit_cost,
        "line_items": line_items,
    }

    ledes_content = generate_ledes_1998b(data)

    ledes_data = LedesData(
        invoice_date=req.invoice_date,
        invoice_number=req.invoice_number,
        client_id=matter.client_id,
        matter_id=matter.matter_id,
        invoice_total=invoice_total,
        line_items=[LineItem(**li) for li in line_items],
    )

    logger.info(f"Structured conversion: {len(line_items)} items, total={invoice_total:.2f}")

    return ConversionResponse(
        filename="structured-input",
        status="success",
        extracted_data=ledes_data,
        ledes_content=ledes_content,
    )


@app.post("/convert/text-to-ledes", response_model=ConversionResponse, tags=["Conversion"])
async def convert_text_to_ledes(req: TextConversionRequest):
    """
    Convert pasted text to LEDES 1998B format.

    Extracts invoice data from raw text using regex patterns.
    Returns both extracted data (for form editing) and LEDES output.
    Matter config is applied if matter_name is provided.
    """
    extracted_data = extract_ledes_data(req.text)

    if not extracted_data.get("line_items"):
        raise HTTPException(
            status_code=400,
            detail="No line items found in text. Expected format: 'description US $amount' or 'description R$ amount'.",
        )

    if req.matter_name:
        matter = matter_store.get(req.matter_name)
        if not matter:
            raise HTTPException(status_code=404, detail=f"Matter not found: {req.matter_name}")
        extracted_data["law_firm_id"] = matter.law_firm_id
        extracted_data["law_firm_name"] = matter.law_firm_name
        extracted_data["client_id"] = matter.client_id
        extracted_data["client_name"] = matter.client_name
        extracted_data["matter_id"] = matter.matter_id
        extracted_data["matter_name"] = matter.matter_name
        extracted_data["client_matter_id"] = matter.client_matter_id
        extracted_data["timekeeper_id"] = matter.timekeeper_id
        extracted_data["timekeeper_name"] = matter.timekeeper_name
        extracted_data["timekeeper_classification"] = matter.timekeeper_classification
        extracted_data["unit_cost"] = matter.unit_cost

    ledes_content = generate_ledes_1998b(extracted_data)

    line_items = [LineItem(**item) for item in extracted_data["line_items"]]
    ledes_data = LedesData(
        invoice_date=extracted_data["invoice_date"],
        invoice_number=extracted_data["invoice_number"],
        client_id=extracted_data["client_id"],
        matter_id=extracted_data["matter_id"],
        invoice_total=extracted_data["invoice_total"],
        line_items=line_items,
    )

    logger.info(f"Text conversion: {len(line_items)} items extracted")

    return ConversionResponse(
        filename="text-input",
        status="success",
        extracted_data=ledes_data,
        ledes_content=ledes_content,
    )


# =============================================================================
# MATTERS CRUD
# =============================================================================

@app.get("/matters", response_model=list[MatterResponse], tags=["Matters"])
async def list_matters():
    """List all saved matters."""
    matters = matter_store.list_all()
    return [MatterResponse(**asdict(m)) for m in matters]


@app.get("/matters/{matter_name:path}", response_model=MatterResponse, tags=["Matters"])
async def get_matter(matter_name: str):
    """Get a specific matter by name."""
    matter = matter_store.get(matter_name)
    if not matter:
        raise HTTPException(status_code=404, detail=f"Matter not found: {matter_name}")
    return MatterResponse(**asdict(matter))


@app.post("/matters", response_model=MatterResponse, status_code=201, tags=["Matters"])
async def create_matter(req: MatterRequest):
    """Create a new matter."""
    existing = matter_store.get(req.matter_name)
    if existing:
        raise HTTPException(status_code=409, detail=f"Matter already exists: {req.matter_name}")
    matter = Matter(
        matter_name=req.matter_name, matter_id=req.matter_id,
        client_matter_id=req.client_matter_id, client_id=req.client_id,
        client_name=req.client_name, law_firm_id=req.law_firm_id,
        law_firm_name=req.law_firm_name, timekeeper_id=req.timekeeper_id,
        timekeeper_name=req.timekeeper_name,
        timekeeper_classification=req.timekeeper_classification,
        unit_cost=req.unit_cost, default_task_code=req.default_task_code,
        default_activity_code=req.default_activity_code,
    )
    created = matter_store.create(matter)
    return MatterResponse(**asdict(created))


@app.put("/matters/{matter_name:path}", response_model=MatterResponse, tags=["Matters"])
async def update_matter(matter_name: str, req: MatterRequest):
    """Update an existing matter."""
    updates = req.model_dump(exclude={"matter_name"})
    updated = matter_store.update(matter_name, updates)
    if not updated:
        raise HTTPException(status_code=404, detail=f"Matter not found: {matter_name}")
    return MatterResponse(**asdict(updated))


@app.delete("/matters/{matter_name:path}", status_code=204, tags=["Matters"])
async def delete_matter(matter_name: str):
    """Delete a matter."""
    if not matter_store.delete(matter_name):
        raise HTTPException(status_code=404, detail=f"Matter not found: {matter_name}")


# =============================================================================
# VALIDATION
# =============================================================================

@app.post("/validate", response_model=ValidationResponse, tags=["Validation"])
async def validate_ledes(ledes_content: str = Form(...)):
    """Validate LEDES 1998B content and return issues."""
    issues = validate_ledes_1998b(ledes_content)
    errors = [i for i in issues if i.severity == "error"]
    warnings = [i for i in issues if i.severity == "warning"]
    return ValidationResponse(
        valid=len(errors) == 0,
        error_count=len(errors),
        warning_count=len(warnings),
        issues=[ValidationIssueResponse(
            line=i.line, field=i.field, field_name=i.field_name,
            severity=i.severity, message=i.message,
        ) for i in issues],
    )


# =============================================================================
# BATCH CONVERSION
# =============================================================================

def _process_single_docx(content: bytes, filename: str, config_dict: dict | None) -> BatchResultItem:
    """Process a single DOCX file for batch conversion."""
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", mode='wb') as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        os.chmod(tmp_path, 0o600)

        try:
            doc = docx.Document(tmp_path)
        except Exception:
            return BatchResultItem(filename=filename, status="error", error="Unable to parse DOCX file")

        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)

        text_content = "\n".join(full_text)
        if not text_content.strip():
            return BatchResultItem(filename=filename, status="error", error="No text content found")

        extracted_data = extract_ledes_data(text_content)

        if not extracted_data.get("line_items"):
            return BatchResultItem(filename=filename, status="error", error="No line items found")

        # Apply config
        if config_dict:
            for key, value in config_dict.items():
                if value is not None:
                    extracted_data[key] = value

        ledes_content = generate_ledes_1998b(extracted_data)

        # Validate
        issues = validate_ledes_1998b(ledes_content)
        validation_issues = [ValidationIssueResponse(
            line=i.line, field=i.field, field_name=i.field_name,
            severity=i.severity, message=i.message,
        ) for i in issues]

        errors = [i for i in issues if i.severity == "error"]

        line_items = [LineItem(**item) for item in extracted_data["line_items"]]
        ledes_data = LedesData(
            invoice_date=extracted_data["invoice_date"],
            invoice_number=extracted_data["invoice_number"],
            client_id=extracted_data["client_id"],
            matter_id=extracted_data["matter_id"],
            invoice_total=extracted_data["invoice_total"],
            line_items=line_items,
        )

        return BatchResultItem(
            filename=filename,
            status="success" if not errors else "warning",
            ledes_content=ledes_content,
            extracted_data=ledes_data,
            validation_issues=validation_issues,
        )
    except Exception as e:
        logger.error(f"Batch processing failed for {filename}: {e}")
        return BatchResultItem(filename=filename, status="error", error=str(e))
    finally:
        if 'tmp_path' in locals() and os.path.exists(tmp_path):
            os.remove(tmp_path)


@app.post("/convert/batch", response_model=BatchConversionResponse, tags=["Conversion"])
async def convert_batch(
    request: Request,
    files: list[UploadFile] = File(..., description="Multiple DOCX files"),
    matter_name: Annotated[Optional[str], Form(description="Matter name to load config from")] = None,
    config: Annotated[Optional[str], Form(description="JSON config string")] = None,
):
    """Convert multiple DOCX files to LEDES 1998B format."""
    client_ip = request.client.host if request.client else "unknown"
    if not rate_limit_check(client_ip):
        raise HTTPException(status_code=429, detail="Too many requests.")

    # Build config dict from matter or JSON
    config_dict = None
    if matter_name:
        matter = matter_store.get(matter_name)
        if not matter:
            raise HTTPException(status_code=404, detail=f"Matter not found: {matter_name}")
        config_dict = {
            "law_firm_id": matter.law_firm_id,
            "law_firm_name": matter.law_firm_name,
            "client_id": matter.client_id,
            "client_name": matter.client_name,
            "matter_id": matter.matter_id,
            "client_matter_id": matter.client_matter_id,
            "timekeeper_id": matter.timekeeper_id,
            "timekeeper_name": matter.timekeeper_name,
            "timekeeper_classification": matter.timekeeper_classification,
            "unit_cost": matter.unit_cost,
        }
    elif config:
        try:
            config_dict = json.loads(config)
        except json.JSONDecodeError as e:
            raise HTTPException(status_code=400, detail=f"Invalid JSON: {e}")

    results: list[BatchResultItem] = []
    for file in files:
        if not file.filename or not file.filename.lower().endswith('.docx'):
            results.append(BatchResultItem(
                filename=file.filename or "unknown",
                status="error",
                error="Invalid file type. Must be .docx",
            ))
            continue

        content = await file.read()
        if len(content) == 0:
            results.append(BatchResultItem(
                filename=file.filename, status="error", error="Empty file",
            ))
            continue

        if len(content) > MAX_FILE_SIZE:
            results.append(BatchResultItem(
                filename=file.filename, status="error", error="File too large",
            ))
            continue

        if not validate_file_type(content, file.filename):
            results.append(BatchResultItem(
                filename=file.filename, status="error", error="Invalid DOCX format",
            ))
            continue

        result = _process_single_docx(content, file.filename, config_dict)
        results.append(result)

    successful = sum(1 for r in results if r.status == "success")
    failed = sum(1 for r in results if r.status == "error")

    return BatchConversionResponse(
        total=len(results),
        successful=successful,
        failed=failed,
        results=results,
    )


@app.get("/debug/sentry", tags=["Debug"])
async def debug_sentry():
    """Test Sentry integration."""
    raise Exception("Sentry test exception from ledes-converter")
