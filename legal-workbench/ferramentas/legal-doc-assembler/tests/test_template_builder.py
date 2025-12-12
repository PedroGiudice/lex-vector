# tests/test_template_builder.py
import pytest
import tempfile
import json
from pathlib import Path
from docx import Document
from src.template_builder import TemplateBuilder

@pytest.fixture
def sample_docx():
    """Create a sample DOCX for testing."""
    doc = Document()
    doc.add_paragraph("Cliente: João da Silva")
    doc.add_paragraph("CPF: 123.456.789-01")
    doc.add_paragraph("Valor: R$ 1.234,56")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        yield f.name
    Path(f.name).unlink(missing_ok=True)

@pytest.fixture
def templates_dir():
    """Create temp templates directory."""
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)

def test_builder_loads_docx(sample_docx):
    """Builder should load and parse DOCX file."""
    builder = TemplateBuilder(sample_docx)
    assert builder.get_full_text() is not None
    assert "João da Silva" in builder.get_full_text()

def test_builder_detects_patterns(sample_docx):
    """Builder should detect patterns automatically."""
    builder = TemplateBuilder(sample_docx)
    detections = builder.detect_patterns()

    assert len(detections) >= 2  # CPF and valor
    types = [d['type'] for d in detections]
    assert 'cpf' in types
    assert 'valor' in types

def test_builder_replaces_selection(sample_docx):
    """Builder should replace selected text with Jinja2 variable."""
    builder = TemplateBuilder(sample_docx)

    # Replace "João da Silva" with {{ nome }}
    builder.add_field_replacement(
        text="João da Silva",
        field_name="nome",
        filter_name="nome"
    )

    result = builder.get_modified_text()
    assert "{{ nome | nome }}" in result
    assert "João da Silva" not in result

def test_builder_saves_template(sample_docx, templates_dir):
    """Builder should save template DOCX and metadata JSON."""
    builder = TemplateBuilder(sample_docx)
    builder.add_field_replacement("João da Silva", "nome", "nome")
    builder.add_field_replacement("123.456.789-01", "cpf", "cpf")

    result = builder.save_template(
        output_dir=templates_dir,
        template_name="contrato_teste",
        description="Template de teste"
    )

    assert result['success'] is True
    assert (templates_dir / "contrato_teste.docx").exists()
    assert (templates_dir / "contrato_teste.json").exists()

    # Check metadata
    with open(templates_dir / "contrato_teste.json") as f:
        meta = json.load(f)
    assert meta['name'] == "contrato_teste"
    assert len(meta['fields']) == 2
    assert 'nome' in [f['name'] for f in meta['fields']]
