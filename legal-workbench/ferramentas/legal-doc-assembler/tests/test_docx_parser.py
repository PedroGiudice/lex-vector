# tests/test_docx_parser.py
import pytest
from pathlib import Path
from src.docx_parser import DocxParser

def test_extract_paragraphs_returns_list():
    """Parser should extract all paragraphs from a DOCX file."""
    # Create a temp DOCX for testing
    from docx import Document
    import tempfile

    doc = Document()
    doc.add_paragraph("First paragraph")
    doc.add_paragraph("Second paragraph")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)

        parser = DocxParser(f.name)
        paragraphs = parser.extract_paragraphs()

        assert isinstance(paragraphs, list)
        assert len(paragraphs) == 2
        assert paragraphs[0]['text'] == "First paragraph"
        assert paragraphs[1]['text'] == "Second paragraph"
