# tests/test_integration.py
"""Integration tests for the complete template creation workflow."""

import pytest
import tempfile
import json
from pathlib import Path
from docx import Document

from src.template_builder import TemplateBuilder
from src.template_manager import TemplateManager
from src.engine import DocumentEngine


@pytest.fixture
def sample_contract():
    """Create a realistic sample contract DOCX."""
    doc = Document()
    doc.add_paragraph("CONTRATO DE PRESTAÇÃO DE SERVIÇOS")
    doc.add_paragraph("")
    doc.add_paragraph("CONTRATANTE: Maria das Graças Silva")
    doc.add_paragraph("CPF: 123.456.789-01")
    doc.add_paragraph("Endereço: Rua das Flores, 42, São Paulo/SP")
    doc.add_paragraph("CEP: 01310-100")
    doc.add_paragraph("")
    doc.add_paragraph("CONTRATADA: Empresa XYZ Ltda")
    doc.add_paragraph("CNPJ: 12.345.678/0001-99")
    doc.add_paragraph("")
    doc.add_paragraph("Valor: R$ 5.000,00 (cinco mil reais)")

    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as f:
        doc.save(f.name)
        yield f.name
    Path(f.name).unlink(missing_ok=True)


@pytest.fixture
def templates_dir():
    with tempfile.TemporaryDirectory() as d:
        yield Path(d)


class TestFullWorkflow:
    """Test complete workflow: create template -> save -> use."""

    def test_create_and_use_template(self, sample_contract, templates_dir):
        """Full workflow: create template, save it, then use it."""

        # 1. Create template from sample contract
        builder = TemplateBuilder(sample_contract)

        # 2. Detect patterns
        patterns = builder.detect_patterns()
        assert len(patterns) >= 4  # CPF, CNPJ, CEP, valor

        # 3. Apply detected patterns
        for pattern in patterns:
            builder.apply_detected_pattern(pattern)

        # 4. Add manual field for name
        builder.add_field_replacement(
            "Maria das Graças Silva",
            "nome_contratante",
            "nome"
        )

        # 5. Save template
        result = builder.save_template(
            output_dir=templates_dir,
            template_name="Contrato de Serviços",
            description="Template para contratos de prestação de serviços",
            tags=["contrato", "serviços"]
        )
        assert result['success'] is True

        # 6. Verify template saved
        manager = TemplateManager(templates_dir)
        templates = manager.list_templates()
        assert len(templates) == 1
        assert templates[0]['name'] == "Contrato de Serviços"

        # 7. Use template with new data
        engine = DocumentEngine()
        template_path = Path(result['docx_path'])  # Use actual saved path

        new_data = {
            "nome_contratante": "João Carlos Santos",
            "cpf": "98765432100",
            "cep": "04567890",
            "cnpj": "99888777000166",
            "valor": "10000.00"
        }

        # Validate data against template
        validation = engine.validate_data(str(template_path), new_data)

        # Should have no missing required fields
        # (Note: we may have extra fields that aren't in template)
        assert 'nome_contratante' not in validation.get('missing', [])

        # 8. Render document
        output_path = templates_dir / "output.docx"
        engine.render(
            str(template_path),
            new_data,
            str(output_path)
        )

        assert output_path.exists()

        # 9. Verify output content
        output_doc = Document(str(output_path))
        output_text = '\n'.join(p.text for p in output_doc.paragraphs)

        # Check normalized values appear
        assert "João Carlos Santos" in output_text or "JOÃO CARLOS SANTOS" in output_text
