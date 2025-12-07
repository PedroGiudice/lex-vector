#!/usr/bin/env python3
"""
Script to generate test template for POC.
Creates a .docx file with test placeholders programmatically.
"""

import json
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from rich.console import Console

console = Console()


def create_test_template() -> Path:
    """Create a comprehensive test template with various field types."""

    output_dir = Path(__file__).parent / "templates"
    output_dir.mkdir(exist_ok=True)
    output_path = output_dir / "template_teste.docx"

    doc = Document()

    # Title
    title = doc.add_heading("DOCUMENTO DE TESTE", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Personal data section
    doc.add_heading("1. DADOS PESSOAIS", level=2)

    p1 = doc.add_paragraph()
    p1.add_run("Nome: ").bold = True
    p1.add_run("{{ nome }}")

    p2 = doc.add_paragraph()
    p2.add_run("CPF: ").bold = True
    p2.add_run("{{ cpf }}")

    p3 = doc.add_paragraph()
    p3.add_run("RG: ").bold = True
    p3.add_run("{{ rg }}")

    p4 = doc.add_paragraph()
    p4.add_run("Nacionalidade: ").bold = True
    p4.add_run("{{ nacionalidade }}")

    p5 = doc.add_paragraph()
    p5.add_run("Estado Civil: ").bold = True
    p5.add_run("{{ estado_civil }}")

    p6 = doc.add_paragraph()
    p6.add_run("Profissão: ").bold = True
    p6.add_run("{{ profissao }}")

    doc.add_paragraph()

    # Address section
    doc.add_heading("2. ENDEREÇO", level=2)

    p7 = doc.add_paragraph()
    p7.add_run("Logradouro: ").bold = True
    p7.add_run("{{ endereco }}")

    p8 = doc.add_paragraph()
    p8.add_run("Número: ").bold = True
    p8.add_run("{{ numero }}")

    p9 = doc.add_paragraph()
    p9.add_run("Complemento: ").bold = True
    p9.add_run("{{ complemento }}")

    p10 = doc.add_paragraph()
    p10.add_run("Bairro: ").bold = True
    p10.add_run("{{ bairro }}")

    p11 = doc.add_paragraph()
    p11.add_run("Cidade/UF: ").bold = True
    p11.add_run("{{ cidade }}/{{ uf }}")

    p12 = doc.add_paragraph()
    p12.add_run("CEP: ").bold = True
    p12.add_run("{{ cep }}")

    doc.add_paragraph()

    # Legal section (with filters example)
    doc.add_heading("3. DADOS COM FILTROS", level=2)

    doc.add_paragraph(
        "Este documento demonstra o uso de filtros para normalização automática:"
    )

    p13 = doc.add_paragraph()
    p13.add_run("Nome (filtro nome): ").bold = True
    p13.add_run("{{ nome_raw|nome }}")

    p14 = doc.add_paragraph()
    p14.add_run("Endereço (filtro endereco): ").bold = True
    p14.add_run("{{ endereco_raw|endereco }}")

    p15 = doc.add_paragraph()
    p15.add_run("CPF (filtro cpf): ").bold = True
    p15.add_run("{{ cpf_raw|cpf }}")

    p16 = doc.add_paragraph()
    p16.add_run("CNPJ (filtro cnpj): ").bold = True
    p16.add_run("{{ cnpj_raw|cnpj }}")

    p17 = doc.add_paragraph()
    p17.add_run("CEP (filtro cep): ").bold = True
    p17.add_run("{{ cep_raw|cep }}")

    doc.add_paragraph()

    # Signature section
    doc.add_heading("4. ASSINATURA", level=2)

    doc.add_paragraph()
    doc.add_paragraph()

    sig = doc.add_paragraph("_" * 50)
    sig.alignment = WD_ALIGN_PARAGRAPH.CENTER

    name_sig = doc.add_paragraph("{{ nome }}")
    name_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cpf_sig = doc.add_paragraph("CPF: {{ cpf }}")
    cpf_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    date_p = doc.add_paragraph("{{ cidade }}, {{ data }}")
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    # Save
    doc.save(output_path)
    console.print(f"[green]✔[/green] Template criado: {output_path}")

    return output_path


def create_sample_json() -> Path:
    """Create sample JSON data for testing."""

    output_dir = Path(__file__).parent
    output_path = output_dir / "dados_teste.json"

    data = {
        # Dados já formatados
        "nome": "João da Silva",
        "cpf": "123.456.789-01",
        "nacionalidade": "brasileiro",
        "estado_civil": "casado",
        "profissao": "engenheiro",
        "endereco": "Rua das Flores",
        "numero": "123",
        "complemento": "Apto. 101",
        "bairro": "Centro",
        "cidade": "São Paulo",
        "uf": "SP",
        "cep": "01310-100",
        "data": "06 de dezembro de 2025",

        # Dados RAW para teste de filtros
        "nome_raw": "MARIA  DAS GRAÇAS  DA SILVA",
        "endereco_raw": "AV. BRASIL N 500 AP 201 BL A",
        "cpf_raw": "98765432100",
        "cnpj_raw": "12345678000199",
        "cep_raw": "04567890",

        # Campo RG intencionalmente OMITIDO para testar fault-tolerance
    }

    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)

    console.print(f"[green]✔[/green] Dados de teste criados: {output_path}")

    return output_path


def main():
    """Run all setup tasks."""
    console.print("\n[bold cyan]Legal Document Assembler - POC Setup[/bold cyan]\n")

    create_test_template()
    create_sample_json()

    console.print("\n[bold green]✔ Setup POC completo![/bold green]\n")


if __name__ == "__main__":
    main()
