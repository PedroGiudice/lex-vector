"""
DOCX Parser for Template Creation.

Extracts structured content from DOCX files for preview and field detection.
"""

from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, field
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table
import re


@dataclass
class ParsedParagraph:
    """Represents a parsed paragraph with metadata."""
    index: int
    text: str
    style: str
    runs: List[Dict[str, Any]] = field(default_factory=list)

    def to_dict(self) -> Dict[str, Any]:
        return {
            'index': self.index,
            'text': self.text,
            'style': self.style,
            'runs': self.runs
        }


class DocxParser:
    """
    Parser for extracting structured content from DOCX files.

    Used in Template Creator to:
    1. Extract paragraphs with metadata
    2. Extract tables with cell content
    3. Preserve formatting information for preview
    """

    def __init__(self, file_path: str | Path):
        self.file_path = Path(file_path)
        if not self.file_path.exists():
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
        self.document = Document(str(self.file_path))

    def extract_paragraphs(self) -> List[Dict[str, Any]]:
        """
        Extract all paragraphs with text and metadata.

        Returns:
            List of dicts with keys: index, text, style, runs
        """
        paragraphs = []
        for idx, para in enumerate(self.document.paragraphs):
            if para.text.strip():  # Skip empty paragraphs
                parsed = ParsedParagraph(
                    index=idx,
                    text=para.text,
                    style=para.style.name if para.style else "Normal",
                    runs=[{
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic
                    } for run in para.runs]
                )
                paragraphs.append(parsed.to_dict())
        return paragraphs

    def extract_tables(self) -> List[Dict[str, Any]]:
        """
        Extract all tables with cell content.

        Returns:
            List of dicts with table data (rows x cols)
        """
        tables = []
        for idx, table in enumerate(self.document.tables):
            table_data = {
                'index': idx,
                'rows': len(table.rows),
                'cols': len(table.columns) if table.rows else 0,
                'cells': []
            }
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    table_data['cells'].append({
                        'row': row_idx,
                        'col': col_idx,
                        'text': cell.text
                    })
            tables.append(table_data)
        return tables

    def get_full_text(self) -> str:
        """Return all text concatenated for preview."""
        texts = [p.text for p in self.document.paragraphs if p.text.strip()]
        return '\n\n'.join(texts)

    def get_structure_summary(self) -> Dict[str, int]:
        """Return counts of document elements."""
        return {
            'paragraphs': len([p for p in self.document.paragraphs if p.text.strip()]),
            'tables': len(self.document.tables),
            'sections': len(self.document.sections)
        }
