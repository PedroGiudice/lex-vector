import os
import sys
import logging

# Add shared module path for logging and Sentry
sys.path.insert(0, '/app')

# Initialize Sentry BEFORE importing FastAPI for proper instrumentation
try:
    from shared.sentry_config import init_sentry
    init_sentry("ledes-converter")
except ImportError:
    pass  # Sentry not available, continue without it

# Configure structured JSON logging
from shared.logging_config import setup_logging
from shared.middleware import RequestIDMiddleware

log_level = getattr(logging, os.getenv("LOG_LEVEL", "INFO").upper(), logging.INFO)
logger = setup_logging("ledes-converter", level=log_level)

from fastapi import FastAPI, UploadFile, File, HTTPException, Request, Form
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse
from typing import Annotated, Optional
import docx
import re
import tempfile
from datetime import datetime, timezone
from collections import defaultdict
import time
import magic
from defusedxml import ElementTree
import json

from .models import LedesData, ConversionResponse, HealthResponse, LineItem, LedesConfig

app = FastAPI(
    title="LEDES Converter API",
    description="Convert DOCX invoices to LEDES 1998B format",
    version="1.0.0"
)

# Request ID middleware for request tracing
app.add_middleware(RequestIDMiddleware)

# CORS middleware - configured for production security
allowed_origins = os.getenv("ALLOWED_ORIGINS", "http://localhost,http://localhost:3000").split(",")
app.add_middleware(
    CORSMiddleware,
    allow_origins=allowed_origins,
    allow_credentials=True,
    allow_methods=["GET", "POST"],
    allow_headers=["Content-Type", "Authorization"],
)

# Log startup
@app.on_event("startup")
async def startup_event():
    logger.info("Service starting", extra={"event": "startup", "version": "1.0.0"})

# Constants
MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
MAX_FILENAME_LENGTH = 255
MAX_DESCRIPTION_LENGTH = 500
MAX_LINE_ITEMS = 1000
ALLOWED_MIME_TYPES = [
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "application/octet-stream"  # Some systems send DOCX as this
]

# Simple rate limiting (production should use Redis)
rate_limit_storage = defaultdict(list)


def rate_limit_check(client_ip: str, max_requests: int = 10, window_seconds: int = 60) -> bool:
    """Simple in-memory rate limiting. Production should use Redis."""
    now = time.time()
    # Clean old entries
    rate_limit_storage[client_ip] = [
        ts for ts in rate_limit_storage[client_ip]
        if now - ts < window_seconds
    ]

    if len(rate_limit_storage[client_ip]) >= max_requests:
        return False

    rate_limit_storage[client_ip].append(now)
    return True


def validate_file_type(content: bytes, filename: str) -> bool:
    """Validate file type using magic bytes (MIME detection)."""
    try:
        mime = magic.from_buffer(content, mime=True)

        # Check MIME type
        if mime not in ALLOWED_MIME_TYPES:
            logger.warning(f"Invalid MIME type detected: {mime} for file extension .docx")
            return False

        # Additional check: DOCX files should start with PK (ZIP header)
        if not content.startswith(b'PK'):
            logger.warning("File does not have valid ZIP/DOCX header")
            return False

        return True
    except Exception as e:
        logger.error(f"File type validation error: {e}")
        return False


def sanitize_string(value: str, max_length: int = MAX_DESCRIPTION_LENGTH) -> str:
    """Sanitize and truncate string values."""
    if not value:
        return ""
    # Remove control characters and excessive whitespace
    cleaned = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', value)
    cleaned = re.sub(r'\s+', ' ', cleaned).strip()
    return cleaned[:max_length]


def sanitize_ledes_field(value: str, max_length: int = MAX_DESCRIPTION_LENGTH) -> str:
    """
    Sanitize field values for LEDES format compliance.

    LEDES 1998B Requirements:
    - ASCII encoding only (non-ASCII chars removed)
    - No pipe characters (field delimiter)
    - No brackets (line terminator chars)
    - Control characters removed
    """
    if not value:
        return ""

    # First apply general sanitization
    cleaned = sanitize_string(value, max_length)

    # Remove pipe and bracket characters (LEDES reserved chars)
    cleaned = cleaned.replace('|', '').replace('[', '').replace(']', '')

    # Ensure ASCII-only by removing non-ASCII characters
    cleaned = cleaned.encode('ascii', errors='ignore').decode('ascii')

    return cleaned


def parse_currency(value_str: str | None) -> float:
    """Parse currency string to float, removing symbols and formatting."""
    if not value_str:
        return 0.0
    # Remove 'US $', '$', ',', ' ' and other non-numeric chars except '.'
    clean_val = re.sub(r'[^\d.]', '', value_str)
    try:
        return float(clean_val)
    except ValueError:
        return 0.0


def format_ledes_currency(amount: float) -> str:
    """
    Format currency for LEDES 1998B compliance.

    LEDES 1998B Requirements:
    - Up to 14 digits before decimal point
    - Exactly 2 digits after decimal point
    - No currency symbols, commas, or spaces
    - Returns empty string if amount exceeds specification limits
    """
    if amount < 0:
        return ""  # LEDES doesn't support negative amounts in standard fields

    # Format with 2 decimal places
    formatted = f"{amount:.2f}"

    # Check if integer part exceeds 14 digits
    integer_part = formatted.split('.')[0]
    if len(integer_part) > 14:
        logger.warning(f"Currency amount {amount} exceeds LEDES 14-digit limit")
        return ""  # Or could raise exception depending on requirements

    return formatted


def format_date_ledes(date_str: str) -> str:
    """Convert date string to LEDES format (YYYYMMDD)."""
    try:
        # Assuming format "Month DD, YYYY" e.g., "December 15, 2025"
        dt = datetime.strptime(date_str, "%B %d, %Y")
        return dt.strftime("%Y%m%d")
    except ValueError:
        return ""


def extract_ledes_data(text: str) -> dict:
    """Extract invoice data from text content with validation."""
    data = {
        "invoice_date": "",
        "invoice_number": "",
        "client_id": "SALESFORCE",  # Placeholder based on template analysis
        "matter_id": "LITIGATION-BRAZIL",  # Placeholder
        "invoice_total": 0.0,
        "line_items": []
    }

    # Regex patterns (case-insensitive for robustness)
    date_pattern = re.compile(r"Date\s*of\s*Issuance:\s*(.*)", re.IGNORECASE)
    invoice_num_pattern = re.compile(r"Invoice\s*#\s*(\d+)", re.IGNORECASE)
    total_pattern = re.compile(r"Total\s*Gross\s*Amount:\s*(?:US\s*)?\$?([\d,]+\.?\d*)", re.IGNORECASE)

    # Line items pattern
    line_item_pattern = re.compile(r"(.*?)\s*US\s*\$([\d,]+)(?:\s|$)")

    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Extract Header Info
        date_match = date_pattern.search(line)
        if date_match:
            raw_date = date_match.group(1).strip()
            data["invoice_date"] = format_date_ledes(sanitize_string(raw_date, 50))

        inv_num_match = invoice_num_pattern.search(line)
        if inv_num_match:
            data["invoice_number"] = sanitize_string(inv_num_match.group(1).strip(), 50)

        total_match = total_pattern.search(line)
        if total_match:
            data["invoice_total"] = parse_currency(total_match.group(1))

        # Extract Line Items (with limit to prevent DoS)
        if len(data["line_items"]) >= MAX_LINE_ITEMS:
            logger.warning(f"Reached max line items limit ({MAX_LINE_ITEMS})")
            break

        if "Total Gross Amount" not in line:
            item_match = line_item_pattern.search(line)
            if item_match:
                desc = sanitize_string(item_match.group(1).strip())
                amount = parse_currency(item_match.group(2))

                # Filter out likely false positives
                if desc and amount > 0:
                    data["line_items"].append({
                        "description": desc,
                        "amount": amount
                    })

    return data


def generate_ledes_1998b(data: dict) -> str:
    """
    Generate LEDES 1998B format from extracted data.

    LEDES 1998B Specification Compliance:
    - Line 1: LEDES1998B[] identifier
    - Line 2: Header with 24 ALL CAPS field names
    - Lines 3+: Data rows with 24 fields each
    - Every line ends with []
    - Pipe-delimited (|)
    - ASCII encoding only
    - Currency: max 14 digits before decimal, 2 after
    - Date: YYYYMMDD format
    - No pipe or bracket characters in field values
    """
    # Line 1: LEDES 1998B identifier
    lines = ["LEDES1998B[]"]

    # Line 2: Header with exactly 24 fields in specification order
    header = (
        "INVOICE_DATE|INVOICE_NUMBER|CLIENT_ID|LAW_FIRM_MATTER_ID|INVOICE_TOTAL|"
        "BILLING_START_DATE|BILLING_END_DATE|INVOICE_DESCRIPTION|LINE_ITEM_NUMBER|"
        "EXP/FEE/INV_ADJ_TYPE|LINE_ITEM_NUMBER_OF_UNITS|LINE_ITEM_ADJUSTMENT_AMOUNT|"
        "LINE_ITEM_TOTAL|LINE_ITEM_DATE|LINE_ITEM_TASK_CODE|LINE_ITEM_EXPENSE_CODE|"
        "LINE_ITEM_ACTIVITY_CODE|TIMEKEEPER_ID|LINE_ITEM_DESCRIPTION|LAW_FIRM_ID|"
        "LINE_ITEM_UNIT_COST|TIMEKEEPER_NAME|TIMEKEEPER_CLASSIFICATION|CLIENT_MATTER_ID[]"
    )
    lines.append(header)

    # Extract config values with defaults
    billing_start = data.get("billing_start_date", "")
    billing_end = data.get("billing_end_date", "")
    timekeeper_id = data.get("timekeeper_id", "")
    timekeeper_name = data.get("timekeeper_name", "")
    timekeeper_class = data.get("timekeeper_classification", "")
    unit_cost = data.get("unit_cost", 0)
    invoice_desc = data.get("invoice_description", "Legal Services")

    # Data rows: Each line item becomes a row with 24 fields
    for i, item in enumerate(data["line_items"], 1):
        # Calculate units if unit_cost is provided
        units = ""
        if unit_cost and unit_cost > 0:
            calculated_units = item['amount'] / unit_cost
            units = f"{calculated_units:.2f}"

        row = [
            sanitize_ledes_field(data["invoice_date"], 8),           # 1. INVOICE_DATE
            sanitize_ledes_field(data["invoice_number"], 50),        # 2. INVOICE_NUMBER
            sanitize_ledes_field(data["client_id"], 100),            # 3. CLIENT_ID
            sanitize_ledes_field(data["matter_id"], 100),            # 4. LAW_FIRM_MATTER_ID
            format_ledes_currency(data['invoice_total']),            # 5. INVOICE_TOTAL
            sanitize_ledes_field(billing_start, 8),                  # 6. BILLING_START_DATE
            sanitize_ledes_field(billing_end, 8),                    # 7. BILLING_END_DATE
            sanitize_ledes_field(invoice_desc, 100),                 # 8. INVOICE_DESCRIPTION
            str(i),                                                  # 9. LINE_ITEM_NUMBER
            "F",                                                     # 10. EXP/FEE/INV_ADJ_TYPE (F=Fee)
            units,                                                   # 11. LINE_ITEM_NUMBER_OF_UNITS
            "",                                                      # 12. LINE_ITEM_ADJUSTMENT_AMOUNT
            format_ledes_currency(item['amount']),                   # 13. LINE_ITEM_TOTAL
            sanitize_ledes_field(data["invoice_date"], 8),           # 14. LINE_ITEM_DATE
            item.get("task_code", ""),                               # 15. LINE_ITEM_TASK_CODE
            "",                                                      # 16. LINE_ITEM_EXPENSE_CODE
            item.get("activity_code", ""),                           # 17. LINE_ITEM_ACTIVITY_CODE
            sanitize_ledes_field(timekeeper_id, 20),                 # 18. TIMEKEEPER_ID
            sanitize_ledes_field(item["description"], 500),          # 19. LINE_ITEM_DESCRIPTION
            sanitize_ledes_field(data.get("law_firm_id", ""), 50),   # 20. LAW_FIRM_ID
            format_ledes_currency(unit_cost) if unit_cost else "",   # 21. LINE_ITEM_UNIT_COST
            sanitize_ledes_field(timekeeper_name, 50),               # 22. TIMEKEEPER_NAME
            sanitize_ledes_field(timekeeper_class, 10),              # 23. TIMEKEEPER_CLASSIFICATION
            sanitize_ledes_field(data.get("client_matter_id", ""), 50) # 24. CLIENT_MATTER_ID
        ]

        lines.append("|".join(row) + "[]")

    return "\n".join(lines)


@app.get("/health", response_model=HealthResponse)
async def health_check():
    """Health check endpoint for Docker/Kubernetes."""
    return HealthResponse(status="ok", service="ledes-converter")


@app.post("/convert/docx-to-ledes", response_model=ConversionResponse)
async def convert_docx_to_ledes(
    request: Request,
    file: Annotated[UploadFile, File(description="DOCX file to convert to LEDES")],
    config: Annotated[Optional[str], Form(description="JSON string with LEDES configuration")] = None
) -> ConversionResponse:
    """
    Convert a DOCX invoice file to LEDES 1998B format.

    - **file**: DOCX file containing invoice data (max 10MB)
    - **config**: Optional JSON string with LEDES configuration (law_firm_id, client_id, matter_id, etc.)
    - Returns: LEDES formatted content and extracted data

    Security features:
    - Rate limiting (10 requests/minute per IP)
    - MIME type validation using magic bytes
    - File size validation
    - Input sanitization
    """
    # Rate limiting
    client_ip = request.client.host if request.client else "unknown"
    if not rate_limit_check(client_ip):
        raise HTTPException(
            status_code=429,
            detail="Too many requests. Please try again later."
        )

    # Validate filename
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename is required.")

    if len(file.filename) > MAX_FILENAME_LENGTH:
        raise HTTPException(status_code=400, detail="Filename too long.")

    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(status_code=400, detail="Invalid file type. Please upload a .docx file.")

    # Read and validate file size
    content = await file.read()
    if len(content) == 0:
        raise HTTPException(status_code=400, detail="Empty file uploaded.")

    if len(content) > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=413,
            detail=f"File too large. Maximum size is {MAX_FILE_SIZE // (1024*1024)}MB."
        )

    # Validate MIME type using magic bytes
    if not validate_file_type(content, file.filename):
        raise HTTPException(
            status_code=400,
            detail="Invalid file format. File does not appear to be a valid DOCX document."
        )

    tmp_path = ""
    try:
        # Save temp file with secure permissions
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx", mode='wb') as tmp:
            tmp.write(content)
            tmp_path = tmp.name

        # Set restrictive permissions
        os.chmod(tmp_path, 0o600)

        # Process DOCX with error handling
        try:
            doc = docx.Document(tmp_path)
        except Exception as docx_error:
            logger.error(f"Failed to parse DOCX: {docx_error}")
            raise HTTPException(
                status_code=400,
                detail="Unable to parse DOCX file. File may be corrupted or in an unsupported format."
            )

        full_text = []
        for para in doc.paragraphs:
            if para.text:
                full_text.append(para.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text:
                        full_text.append(cell.text)

        text_content = "\n".join(full_text)

        # Validate extracted content
        if not text_content.strip():
            raise HTTPException(
                status_code=400,
                detail="No text content found in DOCX file."
            )

        # Extract Data
        extracted_data = extract_ledes_data(text_content)

        # Parse and validate config if provided
        ledes_config = None
        if config:
            try:
                config_dict = json.loads(config)
                ledes_config = LedesConfig(**config_dict)
                logger.info(f"Config provided: law_firm={ledes_config.law_firm_id}, client={ledes_config.client_id}, matter={ledes_config.matter_id}")
            except json.JSONDecodeError as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid JSON in config parameter: {str(e)}"
                )
            except Exception as e:
                raise HTTPException(
                    status_code=400,
                    detail=f"Invalid config data: {str(e)}"
                )

        # Merge config with extracted data (config takes precedence)
        if ledes_config:
            # Firm/Client/Matter
            extracted_data["law_firm_id"] = ledes_config.law_firm_id
            extracted_data["law_firm_name"] = ledes_config.law_firm_name
            extracted_data["client_id"] = ledes_config.client_id
            extracted_data["client_name"] = ledes_config.client_name
            extracted_data["matter_id"] = ledes_config.matter_id
            extracted_data["matter_name"] = ledes_config.matter_name
            extracted_data["client_matter_id"] = ledes_config.client_matter_id or ""
            # Timekeeper
            extracted_data["timekeeper_id"] = ledes_config.timekeeper_id or ""
            extracted_data["timekeeper_name"] = ledes_config.timekeeper_name or ""
            extracted_data["timekeeper_classification"] = ledes_config.timekeeper_classification or ""
            extracted_data["unit_cost"] = ledes_config.unit_cost or 0
            # Billing period
            extracted_data["billing_start_date"] = ledes_config.billing_start_date or ""
            extracted_data["billing_end_date"] = ledes_config.billing_end_date or ""
        else:
            # Backward compatibility: use placeholders if no config provided
            extracted_data["law_firm_id"] = ""
            extracted_data["client_matter_id"] = ""
            logger.warning("No config provided, using empty law_firm_id (backward compatibility mode)")

        # Validate extracted data
        if not extracted_data.get("invoice_number"):
            logger.warning("Invoice number not found in document")

        if not extracted_data.get("line_items"):
            raise HTTPException(
                status_code=400,
                detail="No line items found in invoice. Please check document format."
            )

        # Generate LEDES
        ledes_content = generate_ledes_1998b(extracted_data)

        logger.info(f"Successfully converted invoice (lines: {len(extracted_data['line_items'])})")

        # Create validated response using Pydantic models
        line_items = [LineItem(**item) for item in extracted_data["line_items"]]
        ledes_data = LedesData(
            invoice_date=extracted_data["invoice_date"],
            invoice_number=extracted_data["invoice_number"],
            client_id=extracted_data["client_id"],
            matter_id=extracted_data["matter_id"],
            invoice_total=extracted_data["invoice_total"],
            line_items=line_items
        )

        return ConversionResponse(
            filename=sanitize_string(file.filename, MAX_FILENAME_LENGTH),
            status="success",
            extracted_data=ledes_data,
            ledes_content=ledes_content
        )

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Conversion failed: {type(e).__name__}: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=500,
            detail="Conversion failed due to an internal error. Please contact support if the issue persists."
        )
    finally:
        if tmp_path and os.path.exists(tmp_path):
            try:
                os.remove(tmp_path)
            except Exception as cleanup_error:
                logger.error(f"Failed to cleanup temp file: {cleanup_error}")


@app.get("/debug/sentry", tags=["Debug"])
async def debug_sentry():
    """
    Test Sentry integration by triggering a test exception.

    This endpoint is for debugging purposes only.
    In production, this should be disabled or protected.
    """
    raise Exception("Sentry test exception from ledes-converter")
