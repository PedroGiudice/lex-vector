"""
Template Builder API routes.

Endpoints for creating templates from DOCX files with pattern detection and annotation.
"""

import os
import re
import json
import shutil
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Any
from uuid import uuid4

from fastapi import APIRouter, HTTPException, UploadFile, File, status
from fastapi.responses import JSONResponse

# Import builder models
from .builder_models import (
    UploadResponse,
    PatternsRequest,
    PatternsResponse,
    PatternMatch,
    SaveTemplateRequest,
    SaveTemplateResponse,
    TemplateListResponse,
    TemplateBuilderInfo,
    TemplateDetailResponse,
    DuplicateTemplateResponse,
)

# Import python-docx for document manipulation
from docx import Document


# ============================================================================
# Configuration
# ============================================================================

# Storage paths
TEMP_UPLOAD_DIR = Path(os.getenv("TEMP_UPLOAD_DIR", "/tmp/builder_uploads"))

# Use DATA_PATH if set, otherwise fallback to /app/templates/builder
DATA_PATH = os.getenv("DATA_PATH", "/data")
BUILDER_TEMPLATES_DIR = Path(os.getenv("BUILDER_TEMPLATES_DIR", f"{DATA_PATH}/templates/builder"))

# Ensure directories exist
TEMP_UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

# Try to create builder templates directory, fallback to /app if permission denied
try:
    BUILDER_TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
except PermissionError:
    # Fallback to app directory if /data is not writable
    BUILDER_TEMPLATES_DIR = Path("/app/templates/builder")
    BUILDER_TEMPLATES_DIR.mkdir(parents=True, exist_ok=True)
    print(f"Warning: Using fallback templates directory: {BUILDER_TEMPLATES_DIR}")

# Router instance
router = APIRouter()


# ============================================================================
# Pattern Detection Configuration
# ============================================================================

PATTERNS = {
    "cpf": {
        "regex": r'\d{3}\.\d{3}\.\d{3}-\d{2}',
        "description": "CPF (formatted)"
    },
    "cnpj": {
        "regex": r'\d{2}\.\d{3}\.\d{3}/\d{4}-\d{2}',
        "description": "CNPJ (formatted)"
    },
    "date_br": {
        "regex": r'\d{2}/\d{2}/\d{4}',
        "description": "Brazilian date (DD/MM/YYYY)"
    },
    "date_iso": {
        "regex": r'\d{4}-\d{2}-\d{2}',
        "description": "ISO date (YYYY-MM-DD)"
    },
    "currency_br": {
        "regex": r'R\$\s*[\d.,]+',
        "description": "Brazilian currency"
    },
    "phone_br": {
        "regex": r'\(\d{2}\)\s*\d{4,5}-\d{4}',
        "description": "Brazilian phone number"
    },
    "email": {
        "regex": r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}',
        "description": "Email address"
    },
    "cep": {
        "regex": r'\d{5}-\d{3}',
        "description": "CEP (postal code)"
    },
    "oab": {
        "regex": r'OAB/[A-Z]{2}\s*\d+',
        "description": "OAB registration"
    },
}


# ============================================================================
# Helper Functions
# ============================================================================

def detect_patterns(text: str) -> List[PatternMatch]:
    """
    Detect all patterns in text.

    Args:
        text: Text to analyze

    Returns:
        List of PatternMatch objects
    """
    matches = []
    pattern_counters = {}

    for pattern_type, pattern_info in PATTERNS.items():
        regex = pattern_info["regex"]

        for match in re.finditer(regex, text):
            # Count occurrences of each pattern type for unique field naming
            if pattern_type not in pattern_counters:
                pattern_counters[pattern_type] = 0
            pattern_counters[pattern_type] += 1

            # Generate field name with counter if multiple occurrences
            if pattern_counters[pattern_type] == 1:
                suggested_field = pattern_type
            else:
                suggested_field = f"{pattern_type}_{pattern_counters[pattern_type]}"

            matches.append(
                PatternMatch(
                    pattern_type=pattern_type,
                    start=match.start(),
                    end=match.end(),
                    value=match.group(),
                    suggested_field=suggested_field,
                )
            )

    # Sort by position
    matches.sort(key=lambda m: m.start)

    return matches


def extract_text_from_docx(file_path: Path) -> Dict[str, Any]:
    """
    Extract text content and metadata from DOCX file.

    Args:
        file_path: Path to DOCX file

    Returns:
        Dict with text_content, paragraphs, and metadata
    """
    doc = Document(str(file_path))

    # Extract paragraphs
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

    # Full text
    text_content = "\n\n".join(paragraphs)

    # Calculate metadata
    word_count = len(text_content.split())
    char_count = len(text_content)

    metadata = {
        "paragraphs": len(paragraphs),
        "word_count": word_count,
        "char_count": char_count,
        "sections": len(doc.sections),
        "tables": len(doc.tables),
    }

    return {
        "text_content": text_content,
        "paragraphs": paragraphs,
        "metadata": metadata,
    }


def apply_annotations_to_docx(
    source_path: Path,
    output_path: Path,
    annotations: List[Dict[str, Any]],
    text_content: str
) -> List[str]:
    """
    Apply field annotations to DOCX document.

    Args:
        source_path: Path to source DOCX
        output_path: Path to save modified DOCX
        annotations: List of annotation dicts
        text_content: Original text content

    Returns:
        List of field names that were applied
    """
    # Load document
    doc = Document(str(source_path))

    # Sort annotations by position (reverse order for safe replacement)
    sorted_annotations = sorted(annotations, key=lambda a: a["start"], reverse=True)

    # Build replacement map
    replacements = {}
    field_names = []

    for annotation in sorted_annotations:
        original = text_content[annotation["start"]:annotation["end"]]
        jinja_var = f"{{{{ {annotation['field_name']} }}}}"
        replacements[original] = jinja_var
        field_names.append(annotation["field_name"])

    # Apply replacements to all paragraphs
    for para in doc.paragraphs:
        for original, jinja_var in replacements.items():
            if original in para.text:
                # Replace in each run
                for run in para.runs:
                    if original in run.text:
                        run.text = run.text.replace(original, jinja_var)

                # Handle text split across runs
                if original in para.text:
                    full_text = para.text.replace(original, jinja_var)
                    # Clear existing runs and add new
                    for run in para.runs:
                        run.text = ""
                    if para.runs:
                        para.runs[0].text = full_text

    # Save modified document
    doc.save(str(output_path))

    return field_names


def save_template_metadata(
    template_id: str,
    template_name: str,
    description: str,
    fields: List[str],
    file_path: str
) -> Dict[str, Any]:
    """
    Save template metadata to JSON file.

    Args:
        template_id: Template UUID
        template_name: Template name
        description: Template description
        fields: List of field names
        file_path: Path to template DOCX

    Returns:
        Metadata dict
    """
    metadata = {
        "id": template_id,
        "name": template_name,
        "description": description,
        "fields": fields,
        "field_count": len(fields),
        "created_at": datetime.utcnow().isoformat() + "Z",
        "file_path": file_path,
    }

    meta_path = BUILDER_TEMPLATES_DIR / f"{template_id}.json"
    with open(meta_path, "w", encoding="utf-8") as f:
        json.dump(metadata, f, indent=2, ensure_ascii=False)

    return metadata


def load_template_metadata(template_id: str) -> Dict[str, Any]:
    """
    Load template metadata from JSON file.

    Args:
        template_id: Template UUID

    Returns:
        Metadata dict

    Raises:
        FileNotFoundError: If metadata file doesn't exist
    """
    meta_path = BUILDER_TEMPLATES_DIR / f"{template_id}.json"
    if not meta_path.exists():
        raise FileNotFoundError(f"Template metadata not found: {template_id}")

    with open(meta_path, "r", encoding="utf-8") as f:
        return json.load(f)


def list_builder_templates() -> List[TemplateBuilderInfo]:
    """
    List all builder templates.

    Returns:
        List of TemplateBuilderInfo objects
    """
    templates = []

    for meta_file in BUILDER_TEMPLATES_DIR.glob("*.json"):
        try:
            with open(meta_file, "r", encoding="utf-8") as f:
                metadata = json.load(f)

            templates.append(
                TemplateBuilderInfo(
                    id=metadata["id"],
                    name=metadata["name"],
                    description=metadata.get("description", ""),
                    field_count=metadata.get("field_count", len(metadata.get("fields", []))),
                    fields=metadata.get("fields", []),
                    created_at=metadata.get("created_at", ""),
                    file_path=metadata.get("file_path", ""),
                )
            )
        except Exception as e:
            # Skip malformed metadata files
            print(f"Warning: Failed to load template metadata from {meta_file}: {e}")
            continue

    # Sort by creation date (newest first)
    templates.sort(key=lambda t: t.created_at, reverse=True)

    return templates


# ============================================================================
# API Endpoints
# ============================================================================

@router.post("/upload", response_model=UploadResponse)
async def upload_document(file: UploadFile = File(...)):
    """
    Upload a DOCX file for template creation.

    Extracts text content and metadata for preview and annotation.

    Args:
        file: Uploaded DOCX file

    Returns:
        UploadResponse with document ID and extracted content
    """
    # Validate file type
    if not file.filename.endswith(".docx"):
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Only .docx files are supported"
        )

    # Generate unique document ID
    document_id = str(uuid4())

    # Save uploaded file
    temp_file_path = TEMP_UPLOAD_DIR / f"{document_id}.docx"

    try:
        with open(temp_file_path, "wb") as f:
            content = await file.read()
            f.write(content)

        # Extract text and metadata
        extracted = extract_text_from_docx(temp_file_path)

        return UploadResponse(
            document_id=document_id,
            text_content=extracted["text_content"],
            paragraphs=extracted["paragraphs"],
            metadata=extracted["metadata"],
        )

    except Exception as e:
        # Clean up on error
        if temp_file_path.exists():
            temp_file_path.unlink()
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to process document: {str(e)}"
        )


@router.post("/patterns", response_model=PatternsResponse)
async def detect_text_patterns(request: PatternsRequest):
    """
    Detect patterns in text (CPF, CNPJ, dates, etc.).

    Analyzes text for common Brazilian legal patterns and suggests field names.

    Args:
        request: PatternsRequest with text to analyze

    Returns:
        PatternsResponse with detected patterns
    """
    try:
        matches = detect_patterns(request.text)

        # Get unique pattern types
        pattern_types_found = list(set(m.pattern_type for m in matches))

        return PatternsResponse(
            matches=matches,
            pattern_types_found=pattern_types_found,
        )

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Pattern detection failed: {str(e)}"
        )


@router.post("/save", response_model=SaveTemplateResponse)
async def save_template(request: SaveTemplateRequest):
    """
    Save template with Jinja2 placeholders.

    Applies field annotations to the uploaded document and saves it as a template.

    Args:
        request: SaveTemplateRequest with document ID and annotations

    Returns:
        SaveTemplateResponse with template information
    """
    # Validate uploaded document exists
    source_path = TEMP_UPLOAD_DIR / f"{request.document_id}.docx"
    if not source_path.exists():
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=f"Uploaded document not found: {request.document_id}"
        )

    try:
        # Extract text content for validation
        extracted = extract_text_from_docx(source_path)
        text_content = extracted["text_content"]

        # Validate annotations are within text bounds
        for annotation in request.annotations:
            if annotation.start < 0 or annotation.end > len(text_content):
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Annotation out of bounds: {annotation.field_name}"
                )

            # Validate original text matches
            actual_text = text_content[annotation.start:annotation.end]
            if actual_text != annotation.original_text:
                raise HTTPException(
                    status_code=status.HTTP_400_BAD_REQUEST,
                    detail=f"Text mismatch for field '{annotation.field_name}': expected '{annotation.original_text}', got '{actual_text}'"
                )

        # Generate template ID
        template_id = str(uuid4())

        # Output path
        output_path = BUILDER_TEMPLATES_DIR / f"{template_id}.docx"

        # Apply annotations
        fields = apply_annotations_to_docx(
            source_path=source_path,
            output_path=output_path,
            annotations=[a.model_dump() for a in request.annotations],
            text_content=text_content,
        )

        # Save metadata
        file_path = f"/templates/builder/{template_id}.docx"
        metadata = save_template_metadata(
            template_id=template_id,
            template_name=request.template_name,
            description=request.description,
            fields=fields,
            file_path=file_path,
        )

        # Clean up uploaded file
        source_path.unlink()

        return SaveTemplateResponse(
            id=template_id,
            name=request.template_name,
            file_path=file_path,
            field_count=len(fields),
            fields=fields,
            created_at=metadata["created_at"],
        )

    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to save template: {str(e)}"
        )


@router.get("/templates", response_model=TemplateListResponse)
async def list_templates():
    """
    List all saved builder templates.

    Returns:
        TemplateListResponse with list of templates
    """
    try:
        templates = list_builder_templates()
        return TemplateListResponse(
            templates=templates,
            count=len(templates),
        )

    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to list templates: {str(e)}"
        )


@router.get("/templates/{template_id}", response_model=TemplateDetailResponse)
async def get_template_detail(template_id: str):
    """
    Get detailed information about a specific template.

    Args:
        template_id: Template UUID

    Returns:
        TemplateDetailResponse with template details
    """
    try:
        metadata = load_template_metadata(template_id)

        template_info = TemplateBuilderInfo(
            id=metadata["id"],
            name=metadata["name"],
            description=metadata.get("description", ""),
            field_count=metadata.get("field_count", len(metadata.get("fields", []))),
            fields=metadata.get("fields", []),
            created_at=metadata.get("created_at", ""),
            file_path=metadata.get("file_path", ""),
        )

        return TemplateDetailResponse(
            template=template_info,
            metadata=metadata,
        )

    except FileNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=str(e)
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to get template details: {str(e)}"
        )


@router.post("/templates/{template_id}/duplicate", response_model=DuplicateTemplateResponse)
async def duplicate_template(template_id: str):
    """
    Duplicate an existing template.

    Creates a copy of the template with a new ID and " (Copy)" appended to the name.

    Args:
        template_id: Template UUID to duplicate

    Returns:
        DuplicateTemplateResponse with new template information
    """
    try:
        # Load original metadata
        original_metadata = load_template_metadata(template_id)

        # Verify template file exists
        original_file = BUILDER_TEMPLATES_DIR / f"{template_id}.docx"
        if not original_file.exists():
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail=f"Template file not found: {template_id}.docx"
            )

        # Generate new template ID
        new_template_id = str(uuid4())

        # Copy DOCX file
        new_file = BUILDER_TEMPLATES_DIR / f"{new_template_id}.docx"
        shutil.copy2(original_file, new_file)

        # Create new metadata
        new_name = f"{original_metadata['name']} (Copy)"
        new_file_path = f"/templates/builder/{new_template_id}.docx"

        new_metadata = save_template_metadata(
            template_id=new_template_id,
            template_name=new_name,
            description=original_metadata.get("description", ""),
            fields=original_metadata.get("fields", []),
            file_path=new_file_path,
        )

        return DuplicateTemplateResponse(
            id=new_template_id,
            name=new_name,
            original_id=template_id,
            file_path=new_file_path,
            created_at=new_metadata["created_at"],
        )

    except FileNotFoundError as e:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail=str(e)
        )
    except Exception as e:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Failed to duplicate template: {str(e)}"
        )
