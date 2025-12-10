#!/usr/bin/env python3
"""
CMR Document Builder - Gerador de DOCX a partir de Markdown
Usa DNA de estilos (JSON) para aplicar formatação consistente.

Uso:
    python doc_builder.py <entrada.md> [saida.docx]
"""

import json
import re
import sys
import os
from typing import Dict, Any, List, Tuple

from docx import Document
from docx.shared import Pt, Twips, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ==========================================
# 1. CARREGAMENTO DE CONFIGURAÇÃO
# ==========================================

def load_dna(json_path: str) -> Dict[str, Any]:
    """Carrega o arquivo de estilos JSON."""
    if not os.path.exists(json_path):
        print(f"Erro: Arquivo de estilos '{json_path}' não encontrado.")
        sys.exit(1)

    with open(json_path, 'r', encoding='utf-8') as f:
        return json.load(f)


# ==========================================
# 2. CONFIGURAÇÃO DE PÁGINA
# ==========================================

def configure_page_setup(doc: Document, dna: Dict[str, Any]) -> None:
    """Aplica margens e tamanho de página."""
    section = doc.sections[0]
    setup = dna.get("page_setup", {})

    # Tamanho
    size = setup.get("size", {})
    if "width_twips" in size:
        section.page_width = Twips(size["width_twips"])
    if "height_twips" in size:
        section.page_height = Twips(size["height_twips"])

    # Margens
    margins = setup.get("margins_twips", {})
    if "top" in margins:
        section.top_margin = Twips(margins["top"])
    if "bottom" in margins:
        section.bottom_margin = Twips(margins["bottom"])
    if "left" in margins:
        section.left_margin = Twips(margins["left"])
    if "right" in margins:
        section.right_margin = Twips(margins["right"])
    if "header" in margins:
        section.header_distance = Twips(margins["header"])
    if "footer" in margins:
        section.footer_distance = Twips(margins["footer"])


# ==========================================
# 3. MOTOR DE ESTILOS
# ==========================================

ALIGN_MAP = {
    "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
    "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
    "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY,
    "BOTH": WD_ALIGN_PARAGRAPH.JUSTIFY
}


def apply_run_formatting(run, style_def: Dict[str, Any]) -> None:
    """Aplica formatação a um run de texto."""
    font = run.font

    # Fonte
    if "font_name" in style_def:
        font.name = style_def["font_name"]
        # Hack XML: Força o Word a reconhecer a fonte em todas as codificações
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), style_def["font_name"])
        rFonts.set(qn('w:hAnsi'), style_def["font_name"])
        rFonts.set(qn('w:eastAsia'), style_def["font_name"])
        rFonts.set(qn('w:cs'), style_def["font_name"])
        rPr.insert(0, rFonts)

    # Tamanho
    if "font_size_pt" in style_def:
        font.size = Pt(style_def["font_size_pt"])

    # Negrito
    if "bold" in style_def:
        font.bold = style_def["bold"]

    # Itálico
    if "italics" in style_def:
        font.italic = style_def["italics"]

    # Cor (padrão: preto)
    if "color" in style_def:
        # Converter hex para RGB
        hex_color = style_def["color"].lstrip('#')
        if len(hex_color) == 6:
            r, g, b = int(hex_color[0:2], 16), int(hex_color[2:4], 16), int(hex_color[4:6], 16)
            font.color.rgb = RGBColor(r, g, b)
    else:
        font.color.rgb = RGBColor(0, 0, 0)


def apply_paragraph_formatting(paragraph, style_def: Dict[str, Any]) -> None:
    """Aplica formatação a um parágrafo."""
    pf = paragraph.paragraph_format

    # Alinhamento
    if "alignment" in style_def:
        pf.alignment = ALIGN_MAP.get(style_def["alignment"], WD_ALIGN_PARAGRAPH.LEFT)

    # Recuos
    if "indent_first_line_twips" in style_def:
        pf.first_line_indent = Twips(style_def["indent_first_line_twips"])
    if "indent_left_twips" in style_def:
        pf.left_indent = Twips(style_def["indent_left_twips"])
    if "indent_hanging_twips" in style_def:
        # Hanging indent é implementado como left_indent + first_line negativo
        pf.left_indent = Twips(style_def.get("indent_left_twips", 0))
        pf.first_line_indent = Twips(-style_def["indent_hanging_twips"])

    # Espaçamento entre linhas
    if "line_spacing_rule" in style_def:
        rule = style_def["line_spacing_rule"]
        if rule == "1.5":
            pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        elif rule == "SINGLE":
            pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
        elif rule == "DOUBLE":
            pf.line_spacing_rule = WD_LINE_SPACING.DOUBLE

    # Espaçamento antes/depois
    if "spacing_before_twips" in style_def:
        pf.space_before = Twips(style_def["spacing_before_twips"])
    if "spacing_after_twips" in style_def:
        pf.space_after = Twips(style_def["spacing_after_twips"])


def add_styled_paragraph(doc: Document, text: str, style_def: Dict[str, Any],
                         inline_formatting: bool = True) -> None:
    """
    Adiciona um parágrafo com estilo aplicado.
    Se inline_formatting=True, processa **negrito** e *itálico*.
    """
    p = doc.add_paragraph()
    apply_paragraph_formatting(p, style_def)

    if inline_formatting:
        # Processar formatação inline
        segments = parse_inline_formatting(text)
        for segment_text, is_bold, is_italic in segments:
            run = p.add_run(segment_text)
            apply_run_formatting(run, style_def)
            if is_bold:
                run.font.bold = True
            if is_italic:
                run.font.italic = True
    else:
        run = p.add_run(text)
        apply_run_formatting(run, style_def)

    return p


# ==========================================
# 4. PARSER DE FORMATAÇÃO INLINE
# ==========================================

def parse_inline_formatting(text: str) -> List[Tuple[str, bool, bool]]:
    """
    Processa formatação inline do Markdown.
    Retorna lista de (texto, is_bold, is_italic).
    """
    segments = []

    # Padrões de formatação (ordem importa: mais específico primeiro)
    patterns = [
        (r'\*\*\*(.+?)\*\*\*', True, True),    # ***bold italic***
        (r'\*\*(.+?)\*\*', True, False),       # **bold**
        (r'\*(.+?)\*', False, True),           # *italic*
        (r'_(.+?)_', False, True),             # _italic_
    ]

    # Encontrar todos os matches
    all_matches = []
    for pattern, is_bold, is_italic in patterns:
        for match in re.finditer(pattern, text):
            all_matches.append({
                'start': match.start(),
                'end': match.end(),
                'text': match.group(1),
                'bold': is_bold,
                'italic': is_italic
            })

    # Ordenar por posição
    all_matches.sort(key=lambda x: x['start'])

    # Remover sobreposições
    filtered = []
    last_end = 0
    for match in all_matches:
        if match['start'] >= last_end:
            filtered.append(match)
            last_end = match['end']

    # Gerar segmentos
    last_idx = 0
    for match in filtered:
        # Texto antes da formatação
        if match['start'] > last_idx:
            plain = text[last_idx:match['start']]
            if plain:
                segments.append((plain, False, False))

        # Texto formatado
        segments.append((match['text'], match['bold'], match['italic']))
        last_idx = match['end']

    # Texto restante
    if last_idx < len(text):
        remaining = text[last_idx:]
        if remaining:
            segments.append((remaining, False, False))

    # Se não encontrou formatação, retorna texto original
    if not segments:
        segments.append((text, False, False))

    return segments


# ==========================================
# 5. HEADER E FOOTER
# ==========================================

def create_header(doc: Document, dna: Dict[str, Any]) -> None:
    """Cria o cabeçalho padrão CMR."""
    header_def = dna.get("header", {})
    if not header_def:
        return

    header = doc.sections[0].header

    # Limpar parágrafos existentes
    for p in list(header.paragraphs):
        p._element.getparent().remove(p._element)

    # Linha 1: "C. M. RODRIGUES"
    line1 = header_def.get("line1", {})
    if line1:
        p1 = header.add_paragraph()
        p1.paragraph_format.alignment = ALIGN_MAP.get(line1.get("alignment", "RIGHT"))
        p1.paragraph_format.space_after = Twips(0)
        run1 = p1.add_run(line1.get("text", ""))
        apply_run_formatting(run1, line1)

    # Linha 2: "Advogados"
    line2 = header_def.get("line2", {})
    if line2:
        p2 = header.add_paragraph()
        p2.paragraph_format.alignment = ALIGN_MAP.get(line2.get("alignment", "RIGHT"))
        run2 = p2.add_run(line2.get("text", ""))
        apply_run_formatting(run2, line2)


def create_footer(doc: Document, dna: Dict[str, Any]) -> None:
    """Cria o rodapé padrão CMR."""
    footer_def = dna.get("footer", {})
    if not footer_def:
        return

    footer = doc.sections[0].footer

    # Limpar parágrafos existentes
    for p in list(footer.paragraphs):
        p._element.getparent().remove(p._element)

    lines = footer_def.get("lines", [])
    for line in lines:
        p = footer.add_paragraph()
        p.paragraph_format.alignment = ALIGN_MAP.get(footer_def.get("alignment", "RIGHT"))
        p.paragraph_format.space_after = Twips(0)
        run = p.add_run(line)
        apply_run_formatting(run, footer_def)


# ==========================================
# 6. PARSER DE MARKDOWN
# ==========================================

def parse_markdown(doc: Document, md_text: str, styles: Dict[str, Dict[str, Any]]) -> None:
    """
    Converte Markdown para elementos DOCX com estilos aplicados.
    """
    lines = md_text.split('\n')
    i = 0

    # Estado para blocos
    in_blockquote = False
    blockquote_lines = []

    def get_style(name: str) -> Dict[str, Any]:
        """Obtém estilo do DNA ou retorna Normal."""
        return styles.get(name, styles.get("Normal", {}))

    def flush_blockquote():
        """Finaliza e adiciona citação acumulada."""
        nonlocal blockquote_lines
        if blockquote_lines:
            quote_text = ' '.join(blockquote_lines)
            style = get_style("Citation")
            add_styled_paragraph(doc, quote_text, style, inline_formatting=False)
            blockquote_lines = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # === LINHA VAZIA ===
        if not stripped:
            if in_blockquote:
                flush_blockquote()
                in_blockquote = False
            i += 1
            continue

        # === LINHA HORIZONTAL (---) ===
        if re.match(r'^-{3,}$', stripped):
            if in_blockquote:
                flush_blockquote()
                in_blockquote = False
            doc.add_page_break()
            i += 1
            continue

        # === TABELA ===
        if stripped.startswith('|') and stripped.endswith('|'):
            if in_blockquote:
                flush_blockquote()
                in_blockquote = False
            # Coletar linhas da tabela
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            parse_table(doc, table_lines, get_style("Normal"))
            continue

        # === BLOCKQUOTE (Citação) ===
        if stripped.startswith('> '):
            in_blockquote = True
            clean = stripped[2:].strip()
            # Remover aspas decorativas
            clean = clean.strip('"').strip("'").strip('"').strip('"')
            blockquote_lines.append(clean)
            i += 1
            continue
        elif in_blockquote:
            # Linha sem > mas ainda parte da citação
            if stripped.startswith('>'):
                blockquote_lines.append(stripped[1:].strip())
                i += 1
                continue
            else:
                flush_blockquote()
                in_blockquote = False

        # === HEADING 1 (# Endereçamento) ===
        if stripped.startswith('# '):
            text = stripped[2:].strip()
            style = get_style("Enderecamento")
            add_styled_paragraph(doc, text, style)
            i += 1
            continue

        # === HEADING 2 (## Título/Capítulo) ===
        if stripped.startswith('## '):
            text = stripped[3:].strip()
            # Heurística: tudo maiúsculo e curto = título da peça
            if text.isupper() and len(text) < 60 and not text.startswith('I'):
                style = get_style("Title_Peca")
            else:
                style = get_style("Chapter_Heading")
            add_styled_paragraph(doc, text, style)
            i += 1
            continue

        # === HEADING 3 (### Subcapítulo) ===
        if stripped.startswith('### '):
            text = stripped[4:].strip()
            style = get_style("Subcapitulo")
            add_styled_paragraph(doc, text, style)
            i += 1
            continue

        # === PARÁGRAFO NUMERADO (1., 2., 3...) ===
        num_match = re.match(r'^(\d+)\.\s+(.+)$', stripped)
        if num_match:
            num = num_match.group(1)
            content = num_match.group(2)
            style = get_style("Paragraph_Numbered")
            p = add_styled_paragraph(doc, f"{num}. {content}", style)
            i += 1
            continue

        # === LISTA COM LETRAS (a), b), c)...) ===
        letter_match = re.match(r'^([a-z]\))\s+(.+)$', stripped, re.IGNORECASE)
        if letter_match:
            marker = letter_match.group(1)
            content = letter_match.group(2)
            style = get_style("Item_Lista")
            # Criar parágrafo com marcador em negrito
            p = doc.add_paragraph()
            apply_paragraph_formatting(p, style)
            run_marker = p.add_run(f"{marker} ")
            apply_run_formatting(run_marker, style)
            run_marker.font.bold = True
            # Conteúdo
            segments = parse_inline_formatting(content)
            for seg_text, is_bold, is_italic in segments:
                run = p.add_run(seg_text)
                apply_run_formatting(run, style)
                if is_bold:
                    run.font.bold = True
                if is_italic:
                    run.font.italic = True
            i += 1
            continue

        # === LISTA COM HÍFEN (- item) ===
        if stripped.startswith('- ') and not re.match(r'^-{3,}$', stripped):
            content = stripped[2:].strip()
            style = get_style("Item_Lista")
            # Verificar se é sub-item romano (i), (ii), etc.
            roman_match = re.match(r'\(([ivx]+)\)\s+(.+)$', content, re.IGNORECASE)
            if roman_match:
                marker = f"({roman_match.group(1)})"
                content = roman_match.group(2)
                p = doc.add_paragraph()
                apply_paragraph_formatting(p, style)
                run_marker = p.add_run(f"{marker} ")
                apply_run_formatting(run_marker, style)
            else:
                p = doc.add_paragraph()
                apply_paragraph_formatting(p, style)
                run_marker = p.add_run("• ")
                apply_run_formatting(run_marker, style)

            segments = parse_inline_formatting(content)
            for seg_text, is_bold, is_italic in segments:
                run = p.add_run(seg_text)
                apply_run_formatting(run, style)
                if is_bold:
                    run.font.bold = True
                if is_italic:
                    run.font.italic = True
            i += 1
            continue

        # === ASSINATURA ===
        if '**' in stripped and 'OAB' in lines[i + 1] if i + 1 < len(lines) else False:
            # Nome em negrito
            name = stripped.replace('**', '').strip()
            style = get_style("Assinatura")
            p = doc.add_paragraph()
            apply_paragraph_formatting(p, style)
            run = p.add_run(name)
            apply_run_formatting(run, style)
            run.font.bold = True
            i += 1
            # Próxima linha: OAB
            if i < len(lines) and 'OAB' in lines[i]:
                oab = lines[i].strip()
                p2 = doc.add_paragraph()
                apply_paragraph_formatting(p2, style)
                p2.paragraph_format.space_before = Twips(0)
                run2 = p2.add_run(oab)
                apply_run_formatting(run2, style)
                i += 1
            continue

        # === TEXTO PADRÃO ===
        style = get_style("Paragraph_Indented")
        add_styled_paragraph(doc, stripped, style)
        i += 1


def parse_table(doc: Document, lines: List[str], style: Dict[str, Any]) -> None:
    """Converte tabela Markdown para tabela DOCX."""
    # Filtrar linhas de separação
    data_lines = [l for l in lines if not re.match(r'^\|[\s\-:]+\|$', l.strip())]

    if not data_lines:
        return

    # Parse das células
    rows = []
    for line in data_lines:
        cells = [c.strip() for c in line.split('|') if c.strip()]
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Criar tabela
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Table Grid'

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx < num_cols:
                cell = table.rows[row_idx].cells[col_idx]
                p = cell.paragraphs[0]
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cell_text)
                apply_run_formatting(run, style)
                if row_idx == 0:
                    run.font.bold = True


# ==========================================
# 7. EXECUÇÃO PRINCIPAL
# ==========================================

def main(input_md: str, output_docx: str) -> None:
    """Função principal de conversão."""
    # Localizar arquivo de estilos
    script_dir = os.path.dirname(os.path.abspath(__file__))
    dna_file = os.path.join(script_dir, 'cmr_styles.json')

    print(f"🧬 Carregando DNA: {dna_file}")
    dna = load_dna(dna_file)

    # Criar documento
    doc = Document()

    # Configurar página
    print("📐 Configurando página...")
    configure_page_setup(doc, dna)

    # Criar header/footer
    print("📝 Criando header/footer...")
    create_header(doc, dna)
    create_footer(doc, dna)

    # Ler markdown
    print(f"📖 Lendo: {input_md}")
    if not os.path.exists(input_md):
        print(f"Erro: Arquivo '{input_md}' não encontrado.")
        sys.exit(1)

    with open(input_md, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # Converter
    print("⚙️  Convertendo markdown...")
    styles = dna.get("styles_map", {})
    parse_markdown(doc, md_content, styles)

    # Salvar
    try:
        doc.save(output_docx)
        print(f"✅ Salvo: {output_docx}")
    except PermissionError:
        print(f"Erro: Não foi possível salvar '{output_docx}'. Arquivo aberto?")
        sys.exit(1)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Uso: python doc_builder.py <entrada.md> [saida.docx]")
        sys.exit(1)

    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else input_file.replace('.md', '.docx')

    main(input_file, output_file)
